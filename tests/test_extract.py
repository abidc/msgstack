"""Unit tests for src/pipeline/extract.py"""

import os
import textwrap
from pathlib import Path

import pytest

os.environ.setdefault("OPENAI_API_KEY", "test-key")
os.environ.setdefault("PINECONE_API_KEY", "test-key")

from src.pipeline.extract import ExtractionError, chunk_text, extract_text


# ── extract_text ─────────────────────────────────────────────────────────────

def test_extract_txt_utf8(tmp_path):
    f = tmp_path / "doc.txt"
    f.write_text("Hello world\nSecond line", encoding="utf-8")
    result = extract_text(f)
    assert "Hello world" in result
    assert "Second line" in result


def test_extract_txt_latin1(tmp_path):
    f = tmp_path / "doc.txt"
    f.write_bytes("Caf\xe9 au lait\n".encode("latin-1"))
    result = extract_text(f)
    assert result.strip() != ""


def test_extract_txt_markdown_extension(tmp_path):
    f = tmp_path / "notes.md"
    f.write_text("# Title\nContent here", encoding="utf-8")
    result = extract_text(f)
    assert "Title" in result


def test_extract_missing_file_raises(tmp_path):
    with pytest.raises(ExtractionError, match="File not found"):
        extract_text(tmp_path / "nonexistent.txt")


def test_extract_unsupported_type_raises(tmp_path):
    f = tmp_path / "data.csv"
    f.write_text("a,b,c", encoding="utf-8")
    with pytest.raises(ExtractionError, match="Unsupported file type"):
        extract_text(f)


def test_extract_docx(tmp_path):
    """DOCX extraction using python-docx."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("First paragraph about the product.")
    doc.add_paragraph("Second paragraph about value.")
    path = tmp_path / "test.docx"
    doc.save(str(path))
    result = extract_text(path)
    assert "First paragraph" in result
    assert "Second paragraph" in result


def test_extract_docx_with_table(tmp_path):
    from docx import Document
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Feature"
    table.cell(0, 1).text = "Benefit"
    table.cell(1, 0).text = "Speed"
    table.cell(1, 1).text = "Saves time"
    path = tmp_path / "table.docx"
    doc.save(str(path))
    result = extract_text(path)
    assert "Feature" in result
    assert "Benefit" in result


def test_extract_pdf(tmp_path):
    """PDF extraction — uses pypdf; test with a minimal PDF."""
    try:
        from pypdf import PdfWriter
    except ImportError:
        pytest.skip("pypdf not available")
    from pypdf import PdfWriter, PageObject
    writer = PdfWriter()
    page = PageObject.create_blank_page(width=200, height=200)
    writer.add_page(page)
    path = tmp_path / "blank.pdf"
    with open(path, "wb") as f:
        writer.write(f)
    # Blank PDF returns empty string — just verify no exception
    result = extract_text(path)
    assert isinstance(result, str)


# ── chunk_text ────────────────────────────────────────────────────────────────

def test_chunk_text_basic():
    text = "A" * 2000
    chunks = chunk_text(text, chunk_size=800, overlap=100)
    assert len(chunks) > 1
    for c in chunks:
        assert "chunk_id" in c
        assert "text" in c
        assert "char_start" in c
        assert "preview" in c


def test_chunk_text_small_doc():
    text = "Short doc"
    chunks = chunk_text(text, chunk_size=800, overlap=100)
    assert len(chunks) == 1
    assert chunks[0]["text"] == text


def test_chunk_text_empty():
    assert chunk_text("") == []
    assert chunk_text("   ") == []


def test_chunk_text_overlap():
    text = "word " * 500  # 2500 chars
    chunks = chunk_text(text, chunk_size=400, overlap=50)
    # Second chunk should start before end of first chunk (overlap)
    assert chunks[1]["char_start"] < chunks[0]["char_end"]


def test_chunk_text_ids_sequential():
    text = "x" * 3000
    chunks = chunk_text(text, chunk_size=800, overlap=0)
    ids = [c["chunk_id"] for c in chunks]
    assert ids == sorted(ids)
