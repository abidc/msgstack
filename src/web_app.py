"""FastAPI web app — admin UX for MsgStack MCP management."""

import logging
import os
import time
import uuid as _uuid
from collections import defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional
from uuid import UUID

from dotenv import load_dotenv
from fastapi import Depends, FastAPI, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse, StreamingResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field, AliasChoices, ConfigDict, model_validator

load_dotenv()

from src.config import settings, estimate_cost_usd
from src.logging_config import configure_logging
configure_logging()

log = logging.getLogger(__name__)

from src.auth import get_auth_context, require_read, require_write, generate_api_key, AuthContext
from src.models import (
    ArtifactStatus, Channel, SchemaType, SpecStatus, AssertionStatus,
    Assertion, Spec, Audience, AssertionType,
)
from src.store import init_store, get_store
from src.pipeline.extract import ExtractionError, extract_text, chunk_text, save_upload
from src.pipeline.structure import SpecStructurer, StructuredSpec
from src.pipeline.skills import SkillManager
from src.rate_limit import extract_limiter, generate_limiter

app = FastAPI(title="MsgStack Admin", version="0.5.0", redirect_slashes=False)
templates = Jinja2Templates(directory="src/web")

app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.cors_origins,
    allow_credentials=settings.cors_origins != ["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/canvas", StaticFiles(directory="src/web/canvas", html=True), name="canvas")

DATA_DIR = Path("data")
DATA_DIR.mkdir(exist_ok=True)
UPLOAD_DIR = DATA_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

store = init_store()

skills = SkillManager(skills_dir=str(DATA_DIR / "skills"))

# Initialize sync engine eagerly so it's available regardless of lifespan routing
from src.sources.sync import init_sync_engine as _init_sync
_sync_engine = _init_sync(store)
structurer = SpecStructurer()

import openai as _oai_mod
_oai_api_key = os.environ.get("OPENAI_API_KEY")
_oai_client = None

def _get_oai_client():
    global _oai_client
    if _oai_client is None:
        if not _oai_api_key:
            raise ValueError("OPENAI_API_KEY environment variable required")
        _oai_client = _oai_mod.OpenAI(api_key=_oai_api_key)
    return _oai_client


@app.on_event("startup")
async def startup_event():
    """Ensure vector index exists on startup, seed if no specs, start sync."""
    from src.grounding.search import GroundingEngine
    try:
        engine = GroundingEngine(
            store=store,
            openai_api_key=os.environ.get("OPENAI_API_KEY"),
            namespace="default",
        )
        engine.ensure_index()
        log.info("Vector index ensured on startup")
    except Exception as e:
        log.warning("Vector index creation skipped: %s", e)

    # Start the background source sync loop (engine already initialized at module level)
    _sync_engine.start()

    # Build the knowledge graph on startup so graph explorer and MCP tools work immediately
    try:
        from src.grounding.graph import get_graph_engine
        get_graph_engine().rebuild()
        log.info("Graph engine rebuilt on startup")
    except Exception as e:
        log.warning("Graph rebuild on startup failed: %s", e)

    # Seed default templates for v0.8 visual system
    try:
        from src.design.template_registry import seed_default_templates
        seed_default_templates()
        log.info("Default templates seeded successfully")
    except Exception as e:
        log.warning("Default templates seeding failed: %s", e)

def _check_token_budget(workspace_id: str) -> None:
    """Raise HTTP 402 if workspace token budget is exhausted."""
    ws = store.get_workspace(workspace_id)
    if not ws or ws.get("max_token_budget", 0) == 0:
        return
    usage = store.get_token_usage_summary(workspace_id=workspace_id)
    used = usage.get("total_input_tokens", 0) + usage.get("total_output_tokens", 0)
    if used >= ws["max_token_budget"]:
        raise HTTPException(
            status_code=402,
            detail=f"Token budget exhausted ({used:,} / {ws['max_token_budget']:,} tokens used). "
                   "Increase max_token_budget via PATCH /api/workspaces/{id}.",
        )


def _now() -> datetime:
    return datetime.now(timezone.utc).replace(tzinfo=None)

# In-memory cache for preview-structure → confirm-structure flow (token → (structured, name, path, ts))
_preview_cache: dict[str, tuple] = {}
_PREVIEW_TTL = 3600  # 1 hour

def _evict_preview_cache() -> None:
    cutoff = time.time() - _PREVIEW_TTL
    stale = [k for k, v in _preview_cache.items() if len(v) > 3 and v[3] < cutoff]
    for k in stale:
        del _preview_cache[k]

# In-memory metrics accumulator
_metrics: dict = defaultdict(lambda: {"requests": 0, "errors": 0, "total_ms": 0.0})
_start_time = time.time()


@app.middleware("http")
async def log_requests(request: Request, call_next):
    start = time.perf_counter()
    response = await call_next(request)
    elapsed_ms = (time.perf_counter() - start) * 1000
    path = request.url.path
    try:
        log.info(
            "%s %s -> %d  (%.1fms)",
            request.method, path, response.status_code, elapsed_ms,
            extra={"endpoint": path, "latency_ms": round(elapsed_ms, 1), "status": response.status_code},
        )
    except Exception:
        pass  # Avoid logging failures
    # Accumulate metrics
    key = f"{request.method} {path}"
    _metrics[key]["requests"] += 1
    _metrics[key]["total_ms"] += elapsed_ms
    if response.status_code >= 400:
        _metrics[key]["errors"] += 1
    return response


# --- Helpers ---

def _spec_response(spec: Spec) -> dict:
    messages = store.get_key_messages(spec.id)
    audiences = store.get_audiences(spec.id)
    completeness = _completeness_score_fast(spec, len(messages), len(audiences))
    entries_list = [
        {
            "id": str(m.id),
            "assertion_type": m.assertion_type.value if hasattr(m.assertion_type, "value") else m.assertion_type,
            "priority": m.priority,
            "content": m.content,
            "variants": m.variants,
            "audiences": m.audiences,
            "channels": [c.value if hasattr(c, "value") else c for c in m.channels],
        }
        for m in messages
    ]
    return {
        "id": str(spec.id),
        "completeness_score": completeness,
        "name": spec.name,
        "source": spec.source,
        "source_id": spec.source_id,
        "status": spec.status,
        "summary": spec.summary,
        "audience": spec.audience,
        "brand_personality": spec.brand_personality,
        "positioning": spec.positioning,
        "tagline": spec.tagline,
        "differentiation": spec.differentiation,
        "last_synced": spec.last_synced.isoformat() if spec.last_synced else None,
        "schema_type": str(spec.schema_type),
        "assertions": entries_list,
        "audiences": [
            {
                "id": str(p.id),
                "name": p.name,
                "description": p.description,
                "qa_pairs": p.qa_pairs,
            }
            for p in audiences
        ],
    }


# --- Specs ---

@app.get("/api/canon-domains")
@app.get("/api/specs")
def list_specs(query: Optional[str] = None, auth: AuthContext = Depends(get_auth_context)):
    workspace_filter = auth.workspace_id if settings.auth_enabled else None
    rows = store.list_specs_with_counts(workspace_id=workspace_filter)
    result = []
    for row in rows:
        h = row["spec"]
        if not auth.has_department_access(h.department):
            continue
        summary = h.summary or ""
        if query and query.lower() not in h.name.lower() and query.lower() not in summary.lower():
            continue
        mc, pc = row["message_count"], row["audience_count"]
        completeness = _completeness_score_fast(h, mc, pc)
        result.append({
            "id": str(h.id),
            "name": h.name,
            "source": h.source,
            "status": h.status,
            "department": h.department,
            "schema_type": str(h.schema_type),
            "summary": summary[:150] + ("..." if len(summary) > 150 else ""),
            "audience_count": pc,
            "message_count": mc,
            "entry_count": mc,
            "last_synced": h.last_synced.isoformat() if h.last_synced else None,
            "completeness_score": completeness,
        })
    return result


def _completeness_score_fast(spec, message_count, audience_count) -> int:
    score = 0
    if spec.name: score += 10
    if spec.summary: score += 10
    if spec.audience: score += 10
    if spec.brand_personality: score += 10
    if spec.positioning: score += 15
    if spec.tagline: score += 10
    if spec.differentiation: score += 10
    if message_count >= 3: score += 10
    if message_count >= 6: score += 10
    if audience_count >= 1: score += 5
    return min(score, 100)


@app.get("/api/canon-domains/{domain_id}")
@app.get("/api/specs/{spec_id}")
def get_spec(spec_id: Optional[str] = None, domain_id: Optional[str] = None, auth: AuthContext = Depends(require_read)):
    actual_id = domain_id or spec_id
    try:
        spec = store.get_spec(UUID(actual_id))
    except Exception:
        raise HTTPException(404, "Invalid ID")
    if not spec:
        raise HTTPException(404, "Spec domain not found")
    if not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to access the '{spec.department}' department.")
    return _spec_response(spec)


# --- Pillars ---

from src.models import PillarCreate, PillarUpdate


@app.get("/api/canon-domains/{domain_id}/pillars")
@app.get("/api/specs/{spec_id}/pillars")
def list_pillars(spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    actual_id = domain_id or spec_id
    try:
        spec_uuid = UUID(actual_id)
    except Exception:
        raise HTTPException(404, "Invalid ID")
    spec = store.get_spec(spec_uuid)
    if not spec:
        raise HTTPException(404, "Spec domain not found")
    pillars = store.list_pillars(spec_uuid)
    return [p.model_dump() for p in pillars]


@app.post("/api/canon-domains/{domain_id}/pillars", status_code=201)
@app.post("/api/specs/{spec_id}/pillars", status_code=201)
def create_pillar(data: PillarCreate, spec_id: Optional[str] = None, domain_id: Optional[str] = None, auth: AuthContext = Depends(require_write)):
    actual_id = domain_id or spec_id
    try:
        spec_uuid = UUID(actual_id)
    except Exception:
        raise HTTPException(404, "Invalid ID")
    spec = store.get_spec(spec_uuid)
    if not spec:
        raise HTTPException(404, "Spec not found")
    if not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to this domain's department.")
    pillar_id = store.create_pillar(spec_uuid, data.name, data.description, data.display_order)
    return {"id": pillar_id, "name": data.name, "description": data.description, "display_order": data.display_order}


@app.patch("/api/canon-domains/{domain_id}/pillars/{pillar_id}")
@app.patch("/api/specs/{spec_id}/pillars/{pillar_id}")
def update_pillar(pillar_id: int, data: PillarUpdate, spec_id: Optional[str] = None, domain_id: Optional[str] = None, auth: AuthContext = Depends(require_write)):
    actual_id = domain_id or spec_id
    try:
        spec_uuid = UUID(actual_id)
    except Exception:
        raise HTTPException(404, "Invalid ID")
    spec = store.get_spec(spec_uuid)
    if not spec:
        raise HTTPException(404, "Spec not found")
    if not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to this domain's department.")
    # Build update dict
    update_dict = {}
    if data.name is not None:
        update_dict["name"] = data.name
    if data.description is not None:
        update_dict["description"] = data.description
    if data.display_order is not None:
        update_dict["display_order"] = data.display_order
    if not update_dict:
        raise HTTPException(400, "No fields to update")
    success = store.update_pillar(pillar_id, **update_dict)
    if not success:
        raise HTTPException(404, "Pillar not found")
    # Return updated pillar
    pillars = store.list_pillars(spec_uuid)
    for p in pillars:
        if p.id == pillar_id:
            return p.model_dump()
    raise HTTPException(404, "Pillar not found")


@app.delete("/api/canon-domains/{domain_id}/pillars/{pillar_id}", status_code=204)
@app.delete("/api/specs/{spec_id}/pillars/{pillar_id}", status_code=204)
def delete_pillar(pillar_id: int, spec_id: Optional[str] = None, domain_id: Optional[str] = None, auth: AuthContext = Depends(require_write)):
    actual_id = domain_id or spec_id
    try:
        spec_uuid = UUID(actual_id)
    except Exception:
        raise HTTPException(404, "Invalid ID")
    spec = store.get_spec(spec_uuid)
    if not spec:
        raise HTTPException(404, "Spec not found")
    if not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to this domain's department.")
    success = store.delete_pillar(pillar_id)
    if not success:
        raise HTTPException(404, "Pillar not found")
    return None


class ChunkPillarAssign(BaseModel):
    pillar_id: int | None = None


@app.patch("/api/chunks/{chunk_id}/pillar")
def assign_chunk_pillar(chunk_id: str, data: ChunkPillarAssign, auth: AuthContext = Depends(require_write)):
    try:
        chunk_uuid = UUID(chunk_id)
    except Exception:
        raise HTTPException(404, "Invalid chunk ID")
    msg = _find_message(chunk_id)
    if not msg:
        raise HTTPException(404, "Chunk not found")
    spec = store.get_spec(msg.spec_id)
    if not spec or not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to this domain's department.")
    success = store.assign_chunk_to_pillar(chunk_uuid, data.pillar_id)
    if not success:
        raise HTTPException(404, "Chunk not found")
    return {"assigned": data.pillar_id}


class SpecCreate(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    name: str
    summary: str = ""
    audience: str = ""
    brand_personality: str = ""
    positioning: str = ""
    tagline: str = ""
    differentiation: str = ""
    source: str = "manual"
    status: str = "active"
    department: str = "General"
    schema_type: SchemaType = SchemaType.ENGINEERING_SPEC


@app.post("/api/canon-domains")
@app.post("/api/specs")
def create_spec(data: SpecCreate, auth: AuthContext = Depends(require_write)):
    if not auth.has_department_access(data.department):
        raise HTTPException(403, f"You do not have permission to write to the '{data.department}' department.")
    
    dept = store.get_department(data.department)
    g_type = data.schema_type
    if g_type == SchemaType.ENGINEERING_SPEC and dept:
        g_type = SchemaType(dept["primary_schema_type"])

    spec = Spec(
        name=data.name,
        source=data.source,
        summary=data.summary,
        audience=data.audience,
        brand_personality=data.brand_personality,
        positioning=data.positioning,
        tagline=data.tagline,
        differentiation=data.differentiation,
        status=SpecStatus(data.status),
        schema_type=g_type,
        department=data.department,
        last_synced=_now(),
    )
    store.upsert_spec(spec)
    return {"id": str(spec.id), "name": spec.name}


class SpecUpdate(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    name: Optional[str] = None
    summary: Optional[str] = None
    audience: Optional[str] = None
    brand_personality: Optional[str] = None
    positioning: Optional[str] = None
    tagline: Optional[str] = None
    differentiation: Optional[str] = None
    status: Optional[str] = None
    department: Optional[str] = None
    schema_type: Optional[SchemaType] = Field(default=None, validation_alias=AliasChoices('schema_type'), serialization_alias='schema_type')


    @model_validator(mode="before")
    @classmethod
    def check_conflicting_types(cls, data):
        if isinstance(data, dict):
            dt = data.get("document_type")
            gt = data.get("schema_type")
            if dt is not None and gt is not None:
                if dt != gt:
                    raise ValueError("Conflicting values provided for 'schema_type' and 'document_type'")
        return data


@app.patch("/api/canon-domains/{domain_id}")
@app.patch("/api/specs/{spec_id}")
def update_spec(data: SpecUpdate, spec_id: Optional[str] = None, domain_id: Optional[str] = None, auth: AuthContext = Depends(require_write)):
    actual_id = domain_id or spec_id
    spec = store.get_spec(UUID(actual_id))
    if not spec:
        raise HTTPException(404, "Spec domain not found")
    if not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to the '{spec.department}' department.")
    if data.department is not None and not auth.has_department_access(data.department):
        raise HTTPException(403, f"You do not have permission to write to the target '{data.department}' department.")

    for k, v in data.model_dump(exclude_none=True).items():
        if k in ("document_type", "schema_type"):
            setattr(spec, "schema_type", SchemaType(v))
        else:
            setattr(spec, k, v)
    store.upsert_spec(spec)
    return {"ok": True, "updated_id": actual_id}


@app.delete("/api/canon-domains/{domain_id}")
@app.delete("/api/specs/{spec_id}")
def delete_spec(spec_id: Optional[str] = None, domain_id: Optional[str] = None, auth: AuthContext = Depends(require_write)):
    actual_id = domain_id or spec_id
    try:
        uid = UUID(actual_id)
    except Exception:
        raise HTTPException(400, "Invalid ID")
    spec = store.get_spec(uid)
    if not spec:
        raise HTTPException(404, "Spec domain not found")
    if not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to the '{spec.department}' department.")
    # Delete vectors BEFORE DB deletion so key message IDs are still available
    try:
        from src.grounding.search import GroundingEngine
        ge = GroundingEngine(
            store=store,
            openai_api_key=os.environ.get("OPENAI_API_KEY"),
            namespace="default",
        )
        ge.delete_spec_vectors(uid)
    except Exception as exc:
        log.warning("Vector index cleanup on domain delete failed: %s", exc)
    store.delete_spec(uid)
    # Rebuild graph so deleted domain is removed immediately
    try:
        from src.grounding.graph import get_graph_engine
        get_graph_engine().rebuild()
    except Exception as exc:
        log.warning("Graph rebuild after domain delete failed: %s", exc)
    return {"ok": True, "deleted_id": actual_id}


# --- Key Messages ---

class EntryCreate(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    spec_id: str = Field(validation_alias=AliasChoices('spec_id', 'canon_domain_id'), serialization_alias='spec_id')
    assertion_type: str
    priority: int = 3
    content: str
    variants: dict = {}
    audiences: list = []
    channels: list = ["all"]
    content_tier: str | None = None




@app.post("/api/entries")
@app.post("/api/messages")
def create_message(data: EntryCreate, auth: AuthContext = Depends(require_write)):
    try:
        domain_id = UUID(data.spec_id)
    except Exception:
        raise HTTPException(400, "Invalid spec ID")
    spec = store.get_spec(domain_id)
    if not spec:
        raise HTTPException(404, "Spec domain not found")
    if not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to this domain's department.")
    try:
        tier = None
        if data.content_tier:
            from src.models import ContentTier
            tier = ContentTier(data.content_tier)
        msg = Assertion(
            spec_id=domain_id,
            assertion_type=AssertionType(data.assertion_type),
            priority=data.priority,
            content=data.content,
            variants=data.variants,
            audiences=data.audiences,
            channels=[Channel(c) for c in data.channels],
            content_tier=tier,
        )
        store.upsert_key_message(msg)
        return {"id": str(msg.id)}
    except Exception as e:
        raise HTTPException(400, str(e))


class EntryUpdate(BaseModel):
    assertion_type: Optional[str] = None
    priority: Optional[int] = None
    content: Optional[str] = None
    variants: Optional[dict] = None
    audiences: Optional[list] = None
    channels: Optional[list] = None

MessageUpdate = EntryUpdate


@app.patch("/api/entries/{entry_id}")
@app.patch("/api/messages/{msg_id}")
def update_message(data: EntryUpdate, msg_id: Optional[str] = None, entry_id: Optional[str] = None, auth: AuthContext = Depends(require_write)):
    actual_id = entry_id or msg_id
    msg = _find_message(actual_id)
    if not msg:
        raise HTTPException(404, "Spec entry not found")
    spec = store.get_spec(msg.spec_id)
    if not spec or not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to this domain's department.")
    if msg.status == AssertionStatus.LOCKED:
        raise HTTPException(409, "Spec entry is locked and cannot be edited. Unlock it first via the status endpoint.")
    for k, v in data.model_dump(exclude_none=True).items():
        if k == "assertion_type":
            msg.assertion_type = AssertionType(v)
        elif k == "channels":
            msg.channels = [Channel(c) for c in v]
        else:
            setattr(msg, k, v)
    store.upsert_key_message(msg)
    return {"ok": True, "updated_id": actual_id}


@app.delete("/api/entries/{entry_id}")
@app.delete("/api/messages/{msg_id}")
def delete_message(msg_id: Optional[str] = None, entry_id: Optional[str] = None, auth: AuthContext = Depends(require_write)):
    actual_id = entry_id or msg_id
    msg = _find_message(actual_id)
    if not msg:
        raise HTTPException(404, "Spec entry not found")
    spec = store.get_spec(msg.spec_id)
    if not spec or not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to this domain's department.")
    if msg.status == AssertionStatus.LOCKED:
        raise HTTPException(409, "Spec entry is locked and cannot be deleted. Unlock it first via the status endpoint.")
    if not store.delete_key_message(UUID(actual_id)):
        raise HTTPException(404, "Spec entry not found")
    return {"ok": True, "deleted_id": actual_id}


# --- Spec Entry Status ---

class EntryStatusUpdate(BaseModel):
    status: str
    approved_by: Optional[str] = None
    notes: Optional[str] = None


@app.patch("/api/entries/{entry_id}/status")
@app.patch("/api/messages/{msg_id}/status")
def update_message_status(data: EntryStatusUpdate, msg_id: Optional[str] = None, entry_id: Optional[str] = None, auth: AuthContext = Depends(require_write)):
    """Update a assertion's approval status."""
    actual_id = entry_id or msg_id
    msg = _find_message(actual_id)
    if not msg:
        raise HTTPException(404, "Spec entry not found")
    spec = store.get_spec(msg.spec_id)
    if not spec or not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to this domain's department.")
    try:
        new_status = AssertionStatus(data.status)
    except ValueError:
        raise HTTPException(400, f"Invalid status: {data.status}")
    # Promotion gate: content_tier required before approving or locking
    if new_status in (AssertionStatus.APPROVED, AssertionStatus.LOCKED) and not msg.content_tier:
        raise HTTPException(400, "Content tier must be assigned before entry can be approved or locked.")
    msg.status = new_status
    if msg.status == AssertionStatus.APPROVED and data.approved_by:
        msg.approved_by = data.approved_by
        msg.approved_at = _now()
    store.upsert_key_message(msg)
    # Log the review action
    store.log_review_action(
        spec_id=msg.spec_id,
        action=f"message_{data.status}",
        performed_by=data.approved_by or "system",
        message_id=UUID(actual_id),
        notes=data.notes or "",
    )
    return {"ok": True, "id": actual_id, "status": str(msg.status)}


class TierUpdate(BaseModel):
    content_tier: str | None


@app.patch("/api/entries/{entry_id}/tier")
@app.patch("/api/messages/{msg_id}/tier")
def update_entry_tier(data: TierUpdate, entry_id: Optional[str] = None, msg_id: Optional[str] = None, auth: AuthContext = Depends(require_write)):
    """Set or clear the content tier on a assertion."""
    actual_id = entry_id or msg_id
    result = store.update_entry_tier(actual_id, data.content_tier)
    if not result:
        raise HTTPException(404, "Spec entry not found")
    return {"ok": True, "id": actual_id, "content_tier": result["content_tier"]}


class DriUpdate(BaseModel):
    dri: str


@app.patch("/api/canon-domains/{domain_id}/dri")
@app.patch("/api/specs/{spec_id}/dri")
def update_domain_dri(data: DriUpdate, domain_id: Optional[str] = None, spec_id: Optional[str] = None, auth: AuthContext = Depends(require_write)):
    """Set the DRI on a spec. Logs a dri_transfer event to the review trail."""
    actual_id = domain_id or spec_id
    result = store.set_domain_dri(actual_id, data.dri, performed_by=auth.name)
    if not result:
        raise HTTPException(404, "Spec domain not found")
    return {"ok": True, "id": actual_id, "dri": result["dri"]}


@app.patch("/api/entries/{entry_id}/dri")
@app.patch("/api/messages/{msg_id}/dri")
def update_entry_dri(data: DriUpdate, entry_id: Optional[str] = None, msg_id: Optional[str] = None, auth: AuthContext = Depends(require_write)):
    """Set the DRI on a assertion (overrides domain DRI). Logs a dri_transfer event."""
    actual_id = entry_id or msg_id
    result = store.set_entry_dri(actual_id, data.dri, performed_by=auth.name)
    if not result:
        raise HTTPException(404, "Spec entry not found")
    return {"ok": True, "id": actual_id, "dri": result["dri"]}


@app.get("/api/dri/summary")
def get_dri_summary():
    """Accountability view: domains grouped by DRI, unowned items first."""
    return store.get_dri_summary()


@app.get("/api/entries/{entry_id}/effective-dri")
@app.get("/api/messages/{msg_id}/effective-dri")
def get_effective_dri_endpoint(entry_id: Optional[str] = None, msg_id: Optional[str] = None):
    """Get the effective DRI for a assertion (entry-level if set, else domain-level)."""
    actual_id = entry_id or msg_id
    dri = store.get_effective_dri(actual_id)
    return {"ok": True, "id": actual_id, "dri": dri}


# --- Spec Review & Staleness ---

@app.post("/api/canon-domains/{domain_id}/review")
@app.post("/api/specs/{spec_id}/review")
def mark_spec_reviewed(spec_id: Optional[str] = None, domain_id: Optional[str] = None, performed_by: Optional[str] = Query(None)):
    """Mark a spec as reviewed (updates last_reviewed timestamp)."""
    actual_id = domain_id or spec_id
    try:
        spec_uuid = UUID(actual_id)
    except Exception:
        raise HTTPException(400, "Invalid ID")
    spec = store.get_spec(spec_uuid)
    if not spec:
        raise HTTPException(404, "Spec domain not found")
    store.update_spec_last_reviewed(spec_uuid)
    store.log_review_action(
        spec_id=spec_uuid,
        action="spec_reviewed",
        performed_by=performed_by or "system",
        notes=f"Spec domain '{spec.name}' marked as reviewed.",
    )
    return {"ok": True, "last_reviewed": _now().isoformat()}


@app.get("/api/canon-domains/{domain_id}/staleness")
@app.get("/api/specs/{spec_id}/staleness")
def check_spec_staleness(spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    """Check if a domain is stale (>90 days since last review)."""
    actual_id = domain_id or spec_id
    try:
        spec_uuid = UUID(actual_id)
    except Exception:
        raise HTTPException(400, "Invalid ID")
    spec = store.get_spec(spec_uuid)
    if not spec:
        raise HTTPException(404, "Spec domain not found")
    is_stale = spec.is_stale(days=90)
    return {
        "domain_id": actual_id,
        "spec_id": actual_id,
        "is_stale": is_stale,
        "last_reviewed": spec.last_reviewed.isoformat() if spec.last_reviewed else None,
        "days_since_review": (
            (datetime.now() - spec.last_reviewed).days if spec.last_reviewed else None
        ),
    }


class DomainAlignmentScoreRequest(BaseModel):
    text: str
    export_format: Optional[str] = None  # "json" or "markdown"


@app.post("/api/domains/{domain_id}/score_alignment")
def api_score_alignment(domain_id: UUID, req: DomainAlignmentScoreRequest):
    # Verify write scope
    store = get_store()
    domain = store.get_spec(domain_id)
    if not domain:
        raise HTTPException(404, f"Domain {domain_id} not found")
        
    try:
        from src.pipeline.alignment import score_alignment, export_report_to_markdown
        report = score_alignment(req.text, domain_id, store)
        
        # Save score on last run (if keeping tracking history)
        # We can update the last generated artifact score in DB if requested
        
        if req.export_format == "markdown":
            return {"markdown": export_report_to_markdown(report)}
            
        return report
    except Exception as e:
        raise HTTPException(500, f"Alignment check failed: {e}")


class AlignmentScoreRequest(BaseModel):
    content: str


@app.post("/api/canon-domains/{domain_id}/score")
@app.post("/api/specs/{spec_id}/score")
def score_alignment_endpoint(req: AlignmentScoreRequest, spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    """Score arbitrary content against the spec."""
    actual_id = domain_id or spec_id
    from src.pipeline.alignment import AlignmentEngine
    try:
        spec_uuid = UUID(actual_id)
    except Exception:
        raise HTTPException(400, "Invalid ID")
        
    engine = AlignmentEngine(store)
    try:
        report = engine.score(spec_uuid, req.content)
        return report.model_dump()
    except Exception as e:
        log.error(f"Alignment scoring error: {e}", exc_info=True)
        raise HTTPException(500, str(e))


@app.get("/api/canon-domains/{domain_id}/review-trail")
@app.get("/api/specs/{spec_id}/review-trail")
def get_spec_review_trail(spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    """Get the full review audit trail for a domain."""
    actual_id = domain_id or spec_id
    try:
        spec_uuid = UUID(actual_id)
    except Exception:
        raise HTTPException(400, "Invalid ID")
    spec = store.get_spec(spec_uuid)
    if not spec:
        raise HTTPException(404, "Spec domain not found")
    trail = store.get_review_trail(spec_uuid)
    return {"domain_id": actual_id, "spec_id": actual_id, "trail": trail, "count": len(trail)}


class MessageReorder(BaseModel):
    ordered_ids: list[str]


@app.patch("/api/canon-domains/{domain_id}/entries/reorder")
@app.patch("/api/specs/{spec_id}/messages/reorder")
def reorder_messages(data: MessageReorder, spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    """Update priority values so entries appear in the given order."""
    actual_id = domain_id or spec_id
    count = 0
    for priority, msg_id in enumerate(data.ordered_ids, start=1):
        msg = _find_message(msg_id)
        if msg and str(msg.spec_id) == actual_id:
            msg.priority = priority
            store.upsert_key_message(msg)
            count += 1
    return {"ok": True, "reordered_count": count}


class BulkMessageImport(BaseModel):
    rows: list[dict]


@app.post("/api/canon-domains/{domain_id}/entries/bulk-import")
@app.post("/api/specs/{spec_id}/messages/bulk-import")
def bulk_import_messages(data: BulkMessageImport, spec_id: Optional[str] = None, domain_id: Optional[str] = None, auth: AuthContext = Depends(require_write)):
    """Import multiple entries at once."""
    actual_id = domain_id or spec_id
    try:
        spec_uuid = UUID(actual_id)
    except Exception:
        raise HTTPException(400, "Invalid ID")
    spec = store.get_spec(spec_uuid)
    if not spec:
        raise HTTPException(404, "Spec domain not found")
    if not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to this domain's department.")

    if not data.rows:
        raise HTTPException(400, "No rows provided")

    created = []
    errors = []
    existing = store.get_key_messages(spec_uuid)
    next_priority = (max((m.priority for m in existing), default=0) + 1)

    for i, row in enumerate(data.rows):
        try:
            content = row.get("content")
            if not content:
                errors.append({"row": i, "error": "Missing required field: content"})
                continue
            if not isinstance(content, str) or not content.strip():
                errors.append({"row": i, "error": "Invalid content: must be non-empty string"})
                continue

            assertion_type = row.get("assertion_type", "headline")
            try:
                st = AssertionType(assertion_type)
            except ValueError:
                st = AssertionType.CAPABILITY

            msg = Assertion(
                spec_id=spec_uuid,
                assertion_type=st,
                priority=row.get("priority", next_priority + i),
                content=content.strip(),
            )
            store.upsert_key_message(msg)
            created.append(str(msg.id))
        except Exception as e:
            errors.append({"row": i, "error": str(e)})

    return {"created": len(created), "errors": errors, "ids": created}


# --- Personas ---

class PersonaCreate(BaseModel):
    spec_id: str
    name: str
    description: str = ""
    qa_pairs: list = []


@app.post("/api/audiences")
def create_audience(data: PersonaCreate, auth: AuthContext = Depends(require_write)):
    try:
        domain_id = UUID(data.spec_id)
    except Exception:
        raise HTTPException(400, "Invalid spec ID")
    spec = store.get_spec(domain_id)
    if not spec:
        raise HTTPException(404, "Spec domain not found")
    if not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to this domain's department.")
    audience = Audience(
        spec_id=domain_id,
        name=data.name,
        description=data.description,
        qa_pairs=data.qa_pairs,
    )
    store.upsert_audience(audience)
    return {"id": str(audience.id)}


class PersonaUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    qa_pairs: Optional[list] = None


@app.patch("/api/audiences/{audience_id}")
def update_audience(audience_id: str, data: PersonaUpdate, auth: AuthContext = Depends(require_write)):
    audience = _find_audience(audience_id)
    if not audience:
        raise HTTPException(404, "Audience not found")
    spec = store.get_spec(audience.spec_id)
    if not spec or not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to this domain's department.")
    for k, v in data.model_dump(exclude_none=True).items():
        setattr(audience, k, v)
    store.upsert_audience(audience)
    return {"ok": True, "updated_id": audience_id}


@app.delete("/api/audiences/{audience_id}")
def delete_audience(audience_id: str, auth: AuthContext = Depends(require_write)):
    audience = _find_audience(audience_id)
    if not audience:
        raise HTTPException(404, "Audience not found")
    spec = store.get_spec(audience.spec_id)
    if not spec or not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to this domain's department.")
    if not store.delete_audience(UUID(audience_id)):
        raise HTTPException(404, "Audience not found")
    return {"ok": True, "deleted_id": audience_id}


# --- Audience Status ---

class PersonaStatusUpdate(BaseModel):
    status: str
    approved_by: Optional[str] = None
    notes: Optional[str] = None


@app.patch("/api/audiences/{audience_id}/status")
def update_audience_status(audience_id: str, data: PersonaStatusUpdate, auth: AuthContext = Depends(require_write)):
    """Update a audience's approval status."""
    audience = _find_audience(audience_id)
    if not audience:
        raise HTTPException(404, "Audience not found")
    spec = store.get_spec(audience.spec_id)
    if not spec or not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to this domain's department.")
    try:
        audience.status = AssertionStatus(data.status)
    except ValueError:
        raise HTTPException(400, f"Invalid status: {data.status}")
    if audience.status == AssertionStatus.APPROVED and data.approved_by:
        audience.approved_by = data.approved_by
        audience.approved_at = _now()
    store.upsert_audience(audience)
    # Log the review action
    store.log_review_action(
        spec_id=audience.spec_id,
        action=f"audience_{data.status}",
        performed_by=data.approved_by or "system",
        notes=data.notes or "",
    )
    return {"ok": True, "id": audience_id, "status": str(audience.status)}


# --- Source Upload & Processing ---

@app.post("/api/upload")
async def upload_source(file: UploadFile = File(...)):
    """Upload a source file (PDF, DOCX, TXT) and extract its text."""
    if not file.filename:
        raise HTTPException(400, "No filename provided")

    file_path = UPLOAD_DIR / file.filename
    save_upload(file.file, file_path)

    try:
        text = extract_text(file_path)
        chunks = chunk_text(text)
    except ExtractionError as e:
        raise HTTPException(400, str(e))

    return {
        "filename": file.filename,
        "char_count": len(text),
        "word_count": len(text.split()),
        "chunk_count": len(chunks),
        "preview": text[:500],
        "extracted_text": text,
    }


@app.post("/api/extract")
async def extract_upload(
    file: UploadFile = File(...),
    source_name: str = Form(""),
    document_type: str = Form("engineering_spec"),
    department: str = Form("General"),
    _rl: None = Depends(extract_limiter),
    auth: AuthContext = Depends(require_write),
):
    """Upload a file, extract text, structure it, and save as a Spec.

    Returns structured error JSON on failure with which stage failed.
    """
    _check_token_budget(auth.workspace_id if settings.auth_enabled else "default")

    if not auth.has_department_access(department):
        raise HTTPException(403, f"You do not have permission to write to the '{department}' department.")

    file_path = UPLOAD_DIR / file.filename
    save_upload(file.file, file_path)

    try:
        text = extract_text(file_path)
    except ExtractionError as e:
        return JSONResponse(status_code=400, content={
            "status": "failed",
            "error": {"stage": "text_extraction", "message": str(e), "detail": str(e)},
        })

    if not source_name:
        source_name = Path(file.filename).stem

    try:
        structured, llm_usage = structurer.structure(text, source_name=source_name, document_type=document_type)
    except Exception as e:
        log.error("LLM structuring failed for %s: %s", file.filename, e)
        return JSONResponse(status_code=500, content={
            "status": "failed",
            "error": {"stage": "llm_structuring", "message": "LLM structuring failed", "detail": str(e)},
        })

    try:
        store.record_token_usage(
            workspace_id=auth.workspace_id if settings.auth_enabled else "default",
            endpoint="extract",
            model="gpt-4o-mini",
            input_tokens=llm_usage["input_tokens"],
            output_tokens=llm_usage["output_tokens"],
            cost_usd=estimate_cost_usd("gpt-4o-mini", llm_usage["input_tokens"], llm_usage["output_tokens"]),
        )
    except Exception:
        pass

    try:
        spec, indexed, markdown = _commit_structured_spec(
            structured,
            file.filename,
            document_type=document_type,
            raw_markdown=text,
            department=department
        )
    except Exception as e:
        log.error("DB commit failed for %s: %s", file.filename, e)
        return JSONResponse(status_code=500, content={
            "status": "failed",
            "error": {"stage": "database_save", "message": "Failed to save to database", "detail": str(e)},
        })

    return {
        "id": str(spec.id),
        "name": spec.name,
        "status": "created",
        "message_count": len(structured.assertions),
        "entry_count": len(structured.assertions),
        "audience_count": len(structured.audiences),
        "indexed": indexed,
        "markdown": markdown,
        "know_your_market": structured.know_your_market,
        "missing_sections": structured.missing_sections,
        "completeness_score": max(0, 100 - len(structured.missing_sections) * 10),
    }


@app.post("/api/preview-structure")
async def preview_structure(
    file: UploadFile = File(...),
    source_name: str = Form(""),
    document_type: str = Form("engineering_spec"),
    department: str = Form("General"),
    _rl: None = Depends(extract_limiter),
    auth: AuthContext = Depends(require_write),
):
    """Extract and structure a file but do NOT save to DB.

    Returns the structured sections for user review. Call /api/confirm-structure
    with the returned preview_token to persist.
    """
    if not auth.has_department_access(department):
        raise HTTPException(403, f"You do not have permission to write to the '{department}' department.")

    file_path = UPLOAD_DIR / file.filename
    save_upload(file.file, file_path)

    try:
        text = extract_text(file_path)
    except ExtractionError as e:
        return JSONResponse(status_code=400, content={
            "status": "failed",
            "error": {"stage": "text_extraction", "message": str(e), "detail": str(e)},
        })

    if not source_name:
        source_name = Path(file.filename).stem

    try:
        structured, llm_usage = structurer.structure(text, source_name=source_name, document_type=document_type)
    except Exception as e:
        log.error("LLM structuring failed for %s: %s", file.filename, e)
        return JSONResponse(status_code=500, content={
            "status": "failed",
            "error": {"stage": "llm_structuring", "message": "LLM structuring failed", "detail": str(e)},
        })

    try:
        store.record_token_usage(
            workspace_id=auth.workspace_id if settings.auth_enabled else "default",
            endpoint="preview-structure",
            model="gpt-4o-mini",
            input_tokens=llm_usage["input_tokens"],
            output_tokens=llm_usage["output_tokens"],
            cost_usd=estimate_cost_usd("gpt-4o-mini", llm_usage["input_tokens"], llm_usage["output_tokens"]),
        )
    except Exception:
        pass

    _evict_preview_cache()
    token = str(_uuid.uuid4())
    _preview_cache[token] = (structured, source_name, str(file_path), time.time(), document_type, department)

    return {
        "status": "preview",
        "preview_token": token,
        "name": structured.name,
        "char_count": len(text),
        "word_count": len(text.split()),
        "summary": structured.summary,
        "audience": structured.audience,
        "brand_personality": structured.brand_personality,
        "positioning": structured.positioning,
        "tagline": structured.tagline,
        "differentiation": structured.differentiation,
        "know_your_market": structured.know_your_market,
        "assertions": structured.assertions,
        "audiences": structured.audiences,
        "missing_sections": structured.missing_sections,
        "completeness_score": max(0, 100 - len(structured.missing_sections) * 10),
    }


@app.post("/api/confirm-structure")
async def confirm_structure(data: dict, auth: AuthContext = Depends(require_write)):
    """Persist a previewed structure to DB and index to Turbovec.

    Body: {"preview_token": "...", "edits": {optional field overrides}}
    """
    token = data.get("preview_token")
    if not token or token not in _preview_cache:
        raise HTTPException(400, "Invalid or expired preview_token")

    cache_entry = _preview_cache.pop(token)
    structured, source_name, file_path_str, timestamp = cache_entry[:4]
    document_type = cache_entry[4] if len(cache_entry) > 4 else "engineering_spec"
    department = cache_entry[5] if len(cache_entry) > 5 else "General"

    # Apply any user edits to top-level fields
    edits = data.get("edits", {})
    for field in ("name", "summary", "audience", "brand_personality", "positioning", "tagline", "differentiation"):
        if field in edits and edits[field]:
            setattr(structured, field, edits[field])
    
    if "document_type" in edits:
        document_type = edits["document_type"]
    if "department" in edits:
        department = edits["department"]

    if not auth.has_department_access(department):
        raise HTTPException(403, f"You do not have permission to write to the '{department}' department.")

    try:
        filename = Path(file_path_str).name
        spec, indexed, markdown = _commit_structured_spec(structured, filename, document_type=document_type, department=department)
    except Exception as e:
        log.error("DB commit failed during confirm-structure: %s", e)
        return JSONResponse(status_code=500, content={
            "status": "failed",
            "error": {"stage": "database_save", "message": "Failed to save to database", "detail": str(e)},
        })

    return {
        "id": str(spec.id),
        "name": spec.name,
        "status": "created",
        "message_count": len(structured.assertions),
        "audience_count": len(structured.audiences),
        "indexed": indexed,
        "markdown": markdown,
        "know_your_market": structured.know_your_market,
        "missing_sections": structured.missing_sections,
        "completeness_score": max(0, 100 - len(structured.missing_sections) * 10),
    }


def _commit_structured_spec(
    structured: StructuredSpec,
    filename: str,
    document_type: str = "engineering_spec",
    raw_markdown: Optional[str] = None,
    department: str = "General"
) -> tuple:
    """Save a StructuredSpec to DB, write markdown, and index to Turbovec.

    Returns (spec, indexed_bool, markdown_str).
    """
    dept = store.get_department(department)
    g_type = SchemaType(document_type)
    if g_type == SchemaType.ENGINEERING_SPEC and dept:
        g_type = SchemaType(dept["primary_schema_type"])

    spec = Spec(
        name=structured.name,
        source="upload",
        source_id=filename,
        schema_type=g_type,
        summary=structured.summary,
        audience=structured.audience,
        brand_personality=structured.brand_personality,
        positioning=structured.positioning,
        tagline=structured.tagline,
        differentiation=structured.differentiation,
        status=SpecStatus.ACTIVE,
        department=department,
        last_synced=_now(),
    )
    store.upsert_spec(spec)

    # ── Phase 2: Create audiences + sub-attrs FIRST (needed for chunk linking) ──
    pain_point_map: dict = {}    # content.lower() → id
    qa_pair_map: dict = {}     # statement.lower() → id

    for p in structured.audiences:
        audience = Audience(
            spec_id=spec.id,
            name=p["name"],
            description=p.get("description", ""),
            qa_pairs=p.get("qa_pairs", []),
        )
        store.upsert_audience(audience)

        audience_obj = store.get_audience_by_name(spec.id, p["name"])
        if audience_obj:
            store.delete_audience_sub_attrs(str(audience_obj.id))

            store.bulk_create_pain_points(
                str(audience_obj.id),
                [pt if isinstance(pt, str) else pt.get("content", str(pt))
                 for pt in p.get("pain_points", [])]
            )
            store.bulk_create_buying_triggers(
                str(audience_obj.id),
                [t if isinstance(t, str) else t.get("content", str(t))
                 for t in p.get("buying_triggers", [])]
            )
            ob_items = []
            for ob in p.get("qa_pairs", []):
                if isinstance(ob, dict):
                    ob_items.append({"statement": ob.get("statement", ""), "response": ob.get("response")})
                else:
                    ob_items.append({"statement": str(ob), "response": None})
            store.bulk_create_qa_pairs(str(audience_obj.id), ob_items)

            for pp in store.list_pain_points(str(audience_obj.id)):
                pain_point_map[pp.content.strip().lower()] = pp.id
            for ob in store.list_qa_pairs(str(audience_obj.id)):
                qa_pair_map[ob.statement.strip().lower()] = ob.id

    # ── Pillars ──────────────────────────────────────────────────────────
    pillar_map = {}  # pillar name -> pillar_id
    for pillar in structured.pillars:
        pillar_id = store.create_pillar(
            spec_id=spec.id,
            name=pillar["name"],
            description=pillar.get("description", ""),
        )
        pillar_map[pillar["name"]] = pillar_id

    # ── Chunks (with Phase 2 linking) ────────────────────────────────
    def _resolve_chunk(chunk_data: dict, pillar_id=None):
        try:
            assertion_type = AssertionType(chunk_data["assertion_type"])
        except ValueError:
            assertion_type = AssertionType.POSITIONING
        msg = Assertion(
            spec_id=spec.id,
            pillar_id=pillar_id,
            assertion_type=assertion_type,
            priority=chunk_data.get("priority", 3),
            content=chunk_data["content"],
            variants=chunk_data.get("variants", {}),
            audiences=chunk_data.get("audiences", []),
            channels=[Channel(c) for c in chunk_data.get("channels", ["all"])],
        )
        store.upsert_key_message(msg)
        # Phase 2: link chunk to pain points / qa_pairs
        pp_ids = [
            pain_point_map[txt.strip().lower()]
            for txt in chunk_data.get("addresses_pain_points", [])
            if txt.strip().lower() in pain_point_map
        ]
        ob_ids = [
            qa_pair_map[txt.strip().lower()]
            for txt in chunk_data.get("resolves_qa_pairs", [])
            if txt.strip().lower() in qa_pair_map
        ]
        if pp_ids or ob_ids:
            # Re-fetch the message to get its UUID
            messages = store.get_key_messages(spec.id)
            for m in messages:
                if m.content == chunk_data["content"] and m.assertion_type == str(assertion_type):
                    store.update_chunk_links(str(m.id), pp_ids, ob_ids)
                    break

    if structured.pillars:
        for pillar in structured.pillars:
            pid = pillar_map.get(pillar["name"])
            if pid is None:
                continue
            for chunk in pillar.get("chunks", []):
                _resolve_chunk(chunk, pillar_id=pid)
        for chunk in structured.ungrouped_chunks:
            _resolve_chunk(chunk)
    else:
        for km in structured.assertions:
            try:
                assertion_type = AssertionType(km["assertion_type"])
            except ValueError:
                assertion_type = AssertionType.POSITIONING
            msg = Assertion(
                spec_id=spec.id,
                pillar_id=None,
                assertion_type=assertion_type,
                priority=km.get("priority", 3),
                content=km["content"],
                variants=km.get("variants", {}),
                audiences=km.get("audiences", []),
                channels=[Channel(c) for c in km.get("channels", ["all"])],
            )
            store.upsert_key_message(msg)

    markdown = structurer.to_markdown(structured)
    save_path = DATA_DIR / "frames" / f"{spec.id}.md"
    save_path.parent.mkdir(exist_ok=True)
    save_path.write_text(markdown, encoding="utf-8")

    if raw_markdown:
        raw_path = DATA_DIR / "sources" / f"{spec.id}.md"
        raw_path.parent.mkdir(exist_ok=True, parents=True)
        raw_path.write_text(raw_markdown, encoding="utf-8")

    from src.grounding.search import GroundingEngine
    spec_row_ws = store.get_spec_workspace_id(spec.id)
    engine = GroundingEngine(
        store=store,
        openai_api_key=os.environ.get("OPENAI_API_KEY"),
        namespace=spec_row_ws or "default",
    )
    try:
        engine.index_spec(spec.id)
        indexed = True
    except Exception as exc:
        log.warning("Vector indexing skipped for %s: %s", spec.id, exc)
        indexed = False

    return spec, indexed, markdown


@app.get("/api/frames/{spec_id}/markdown")
def get_frame_markdown(spec_id: str):
    """Get the saved markdown file for a framework."""
    path = DATA_DIR / "frames" / f"{spec_id}.md"
    if not path.exists():
        raise HTTPException(404, "Markdown file not found")
    return {"markdown": path.read_text()}


# --- Channels ---

class ChannelCreateRequest(BaseModel):
    id: str
    name: str
    description: str = ""


class ChannelUpdateRequest(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None


@app.get("/api/channels")
def get_channels():
    return store.get_channels()


@app.post("/api/channels")
def create_channel(req: ChannelCreateRequest, auth: AuthContext = Depends(require_write)):
    if not req.id or not req.name:
        raise HTTPException(400, "id and name are required")
    if not req.id.replace("_", "").replace("-", "").isalnum():
        raise HTTPException(400, "id must be alphanumeric with underscores/hyphens only")
    return store.upsert_channel(req.id, req.name, req.description)


@app.patch("/api/channels/{channel_id}")
def update_channel(channel_id: str, req: ChannelUpdateRequest,
                   auth: AuthContext = Depends(require_write)):
    channels = store.get_channels()
    existing = next((c for c in channels if c["id"] == channel_id), None)
    if not existing:
        raise HTTPException(404, "Channel not found")
    name = req.name if req.name is not None else existing["name"]
    description = req.description if req.description is not None else existing["description"]
    return store.upsert_channel(channel_id, name, description)


@app.delete("/api/channels/{channel_id}")
def delete_channel(channel_id: str, auth: AuthContext = Depends(require_write)):
    try:
        deleted = store.delete_channel(channel_id)
    except ValueError as e:
        raise HTTPException(400, str(e))
    if not deleted:
        raise HTTPException(404, "Channel not found")
    return {"deleted": True}


class DepartmentCreate(BaseModel):
    name: str
    primary_schema_type: str = "engineering_spec"
    description: str = ""
    workspace_id: str = "default"


@app.get("/api/departments")
def list_departments(auth: AuthContext = Depends(get_auth_context)):
    wid = auth.workspace_id if settings.auth_enabled else "default"
    return store.list_departments(workspace_id=wid)


@app.post("/api/departments")
def create_department(req: DepartmentCreate, auth: AuthContext = Depends(require_write)):
    if not auth.is_admin:
        raise HTTPException(403, "Only administrators can create departments.")
    if not req.name:
        raise HTTPException(400, "name is required")
    from src.models import SchemaType
    try:
        SchemaType(req.primary_schema_type)
    except ValueError:
        raise HTTPException(400, f"Invalid primary grounding type: {req.primary_schema_type}")
    return store.create_department(req.name, req.primary_schema_type, req.description, req.workspace_id)


@app.delete("/api/departments/{name}")
def delete_department(name: str, auth: AuthContext = Depends(require_write)):
    if not auth.is_admin:
        raise HTTPException(403, "Only administrators can delete departments.")
    try:
        deleted = store.delete_department(name)
    except ValueError as e:
        raise HTTPException(400, str(e))
    if not deleted:
        raise HTTPException(404, "Department not found")
    return {"deleted": True}


# --- Knowledge Graph ---

@app.get("/api/graph/stats")
def graph_stats():
    from src.grounding.graph import get_graph_engine
    return get_graph_engine().get_stats()


@app.get("/api/graph/spec/{spec_id}")
def graph_spec(spec_id: str):
    from src.grounding.graph import get_graph_engine
    try:
        UUID(spec_id)
    except ValueError:
        raise HTTPException(400, "Invalid spec_id UUID")
    chunks = get_graph_engine().get_chunks_for_spec(spec_id)
    return {"spec_id": spec_id, "chunks": chunks, "count": len(chunks)}


@app.get("/api/graph/spec/{spec_id}/sections")
def graph_spec_sections(spec_id: str):
    from src.grounding.graph import get_graph_engine
    try:
        UUID(spec_id)
    except ValueError:
        raise HTTPException(400, "Invalid spec_id UUID")
    sections = get_graph_engine().get_sections_for_spec(spec_id)
    return {"spec_id": spec_id, "sections": sections, "count": len(sections)}

@app.get("/api/graph/data")
def graph_data():
    from src.grounding.graph import get_graph_engine
    return get_graph_engine().get_graph_data()


@app.post("/api/graph/rebuild")
def graph_rebuild(auth: AuthContext = Depends(require_write)):
    from src.grounding.graph import get_graph_engine
    engine = get_graph_engine()
    engine.rebuild()
    return {"rebuilt": True, "stats": engine.get_stats()}


# --- Skill Files ---

@app.get("/api/skills")
def list_skills():
    skill_list = skills.list_skills()
    for s in skill_list:
        s["context_inputs"] = skills.get_context_inputs(s["id"])
    return skill_list


@app.get("/api/skills/{skill_id}")
def get_skill(skill_id: str):
    skill = skills.get_skill(skill_id)
    if not skill:
        raise HTTPException(404, f"Skill {skill_id} not found")
    result = dict(skill)
    result["context_inputs"] = skills.get_context_inputs(skill_id)
    return result


@app.put("/api/skills/{skill_id}")
def update_skill(skill_id: str, data: dict):
    return skills.update_skill(skill_id, data)


@app.post("/api/skills")
def create_skill(data: dict):
    if not data.get("id"):
        raise HTTPException(400, "id is required")
    return skills.update_skill(data["id"], data)


@app.delete("/api/skills/{skill_id}")
def delete_skill(skill_id: str):
    if not skills.delete_skill(skill_id):
        raise HTTPException(404, f"Skill {skill_id} not found")
    return {"ok": True}


class ArtifactSectionsUpdate(BaseModel):
    sections: dict


@app.patch("/api/artifacts/{artifact_id}/sections")
def update_artifact_sections(artifact_id: str, data: ArtifactSectionsUpdate):
    """Save manually edited sections back to the artifact record."""
    try:
        aid = UUID(artifact_id)
    except Exception:
        raise HTTPException(400, "Invalid artifact ID")

    with store.session() as s:
        from src.store import ArtifactRecord
        row = s.query(ArtifactRecord).filter(ArtifactRecord.id == aid).first()
        if not row:
            raise HTTPException(404, "Artifact not found")
        
        # Merge sections
        new_sections = row.sections.copy() if row.sections else {}
        new_sections.update(data.sections)
        row.sections = new_sections
        
        # Rebuild raw_content if possible (optional)
        row.raw_content = "\n\n".join([f"### {k}\n{v}" for k, v in new_sections.items()])
        
        s.commit()
    return {"ok": True}


class SectionRegenerate(BaseModel):
    spec_id: str
    skill_id: str
    section_key: str
    context: Optional[dict] = {}


@app.post("/api/generate-section-single")
def regenerate_single_section(data: SectionRegenerate):
    """Regenerate just one part of an artifact using the grounding engine."""
    from src.pipeline.generator import ArtifactGenerator
    generator = ArtifactGenerator(store, skills)
    
    # Custom context to tell the LLM to focus on ONE section
    regen_context = data.context.copy()
    regen_context["focus_section"] = data.section_key
    
    try:
        artifact = generator.generate(data.skill_id, data.spec_id, regen_context)
        # Extract just the section we want
        val = artifact.sections.get(data.section_key, "Regeneration failed to produce this section.")
        return {"section_key": data.section_key, "content": val}
    except Exception as e:
        raise HTTPException(500, str(e))


# --- Stats ---

@app.get("/api/stats")
def get_stats():
    specs = store.list_specs()
    total_messages = sum(len(store.get_key_messages(h.id)) for h in specs)
    total_audiences = sum(len(store.get_audiences(h.id)) for h in specs)
    total_artifacts = sum(len(store.list_artifacts(h.id)) for h in specs)
    usage = store.get_token_usage_summary()

    return {
        "spec_count": len(specs),
        "domain_count": len(specs),
        "message_count": total_messages,
        "entry_count": total_messages,
        "audience_count": total_audiences,
        "artifact_count": total_artifacts,
        "token_count": usage.get("total_input_tokens", 0) + usage.get("total_output_tokens", 0),
        "skill_count": len(skills.list_skills()),
    }


# --- Seed & Index ---

@app.post("/api/seed")
def run_seed():
    """Run the seed script and index all specs to vector index."""
    from seed_data.seed import seed as run_seed_script
    from src.grounding.search import GroundingEngine

    run_seed_script()

    specs = store.list_specs()
    total_messages = sum(len(store.get_key_messages(h.id)) for h in specs)
    total_audiences = sum(len(store.get_audiences(h.id)) for h in specs)

    indexed_count = 0
    for spec in specs:
        ws_id = store.get_spec_workspace_id(spec.id) or "default"
        engine = GroundingEngine(
            store=store,
            openai_api_key=os.environ.get("OPENAI_API_KEY"),
            namespace=ws_id,
        )
        try:
            engine.index_spec(spec.id)
            indexed_count += 1
        except Exception:
            pass

    return {
        "seeded": len(specs),
        "indexed": indexed_count,
        "total_messages": total_messages,
        "total_audiences": total_audiences,
    }


@app.post("/api/canon-domains/{domain_id}/index")
@app.post("/api/specs/{spec_id}/index")
def index_spec(spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    """Index a single domain to vector index."""
    actual_id = domain_id or spec_id
    from src.grounding.search import GroundingEngine

    try:
        spec_uuid = UUID(actual_id)
    except Exception:
        raise HTTPException(400, "Invalid ID")

    spec = store.get_spec(spec_uuid)
    if not spec:
        raise HTTPException(404, "Spec domain not found")

    ws_id = store.get_spec_workspace_id(spec_uuid) or "default"
    engine = GroundingEngine(
        store=store,
        openai_api_key=os.environ.get("OPENAI_API_KEY"),
        namespace=ws_id,
    )

    vectors_indexed = engine.index_spec(spec_uuid)

    return {
        "domain_id": str(spec.id),
        "spec_id": str(spec.id),
        "spec_name": spec.name,
        "vectors_indexed": vectors_indexed,
    }


@app.post("/api/index-all")
def index_all_specs():
    """Index all specs to vector index."""
    from src.grounding.search import GroundingEngine

    specs = store.list_specs()

    total_vectors = 0
    for spec in specs:
        ws_id = store.get_spec_workspace_id(spec.id) or "default"
        engine = GroundingEngine(
            store=store,
            openai_api_key=os.environ.get("OPENAI_API_KEY"),
            namespace=ws_id,
        )
        try:
            vectors = engine.index_spec(spec.id)
            total_vectors += vectors
        except Exception as exc:
            log.warning("Vector index failed for %s: %s", spec.id, exc)

    return {
        "indexed_specs": len(specs),
        "total_vectors": total_vectors,
    }


@app.get("/api/canon-domains/{domain_id}/index-status")
@app.get("/api/specs/{spec_id}/index-status")
def get_spec_index_status(spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    """Check vector index status for a domain: indexed / not_indexed / stale."""
    actual_id = domain_id or spec_id
    try:
        spec_uuid = UUID(actual_id)
    except Exception:
        raise HTTPException(400, "Invalid ID")

    spec = store.get_spec(spec_uuid)
    if not spec:
        raise HTTPException(404, "Spec domain not found")

    try:
        from src.store import VectorMetadataModel
        with store.session() as s:
            record = s.query(VectorMetadataModel).filter(
                VectorMetadataModel.spec_id == str(spec_uuid)
            ).first()

        if not record:
            return {"status": "not_indexed", "spec_id": spec_id}

        # Check staleness: compare last_synced in vector metadata vs spec.last_synced
        meta_dt = record.last_synced
        if meta_dt and spec.last_synced:
            spec_dt = spec.last_synced
            # Make both offset-naive for comparison
            if meta_dt.tzinfo is not None:
                meta_dt = meta_dt.replace(tzinfo=None)
            if spec_dt.tzinfo is not None:
                spec_dt = spec_dt.replace(tzinfo=None)
            stale = meta_dt < spec_dt
            return {
                "status": "stale" if stale else "indexed",
                "spec_id": spec_id,
                "indexed_at": meta_dt.isoformat(),
                "spec_synced": spec.last_synced.isoformat(),
            }

        return {"status": "indexed", "spec_id": spec_id}

    except Exception as e:
        log.error("Index status query failed for %s: %s", spec_id, e)
        return {"status": "error", "message": str(e)}


# --- Internal helpers ---

def _find_message(msg_id: str) -> Optional[Assertion]:
    try:
        return store.get_key_message(UUID(msg_id))
    except (ValueError, Exception):
        return None


def _find_audience(audience_id: str) -> Optional[Audience]:
    try:
        return store.get_audience(UUID(audience_id))
    except (ValueError, Exception):
        return None


# --- Artifact Generation & Preview ---

@app.post("/api/generate")
def generate_artifact(
    skill_id: str = Form(...),
    spec_id: str = Form(...),
    extra_context: Optional[str] = Form(None),
    _rl: None = Depends(generate_limiter),
    auth: AuthContext = Depends(require_read),
):
    """Generate an artifact using a skill and return content for preview."""
    import json as _json
    from src.pipeline.generator import ArtifactGenerator

    context: dict = {}
    if extra_context:
        try:
            context = _json.loads(extra_context)
        except Exception:
            pass

    generator = ArtifactGenerator(store, skills)

    try:
        artifact = generator.generate(skill_id, spec_id, context)
        skill_config = skills.get_skill(skill_id)
        artifact_type = skill_config.get("prefab_template", skill_id) if skill_config else skill_id
        visual_url = None

        if artifact.renderer_type not in ("fabric", "reveal"):
            visual_types = {"one_pager", "social_posts", "email_template", "battlecard", "email_sequence"}
            if artifact_type in visual_types:
                visual_url = f"{settings.base_url}/artifact/{artifact_type}/{spec_id}"
                if skill_id == "battlecard" and context.get("competitor"):
                    visual_url += f"?competitor={context['competitor']}"
                elif skill_id == "email_template" and context.get("stage"):
                    visual_url += f"?stage={context['stage']}"

        # Record token usage
        if artifact.input_tokens or artifact.output_tokens:
            try:
                store.record_token_usage(
                    workspace_id=auth.workspace_id if settings.auth_enabled else "default",
                    endpoint="generate",
                    model="gpt-4o-mini",
                    input_tokens=artifact.input_tokens,
                    output_tokens=artifact.output_tokens,
                    cost_usd=estimate_cost_usd("gpt-4o-mini", artifact.input_tokens, artifact.output_tokens),
                )
            except Exception:
                pass

        # Auto-save to artifact history
        try:
            from src.pipeline.alignment import AlignmentEngine
            score = AlignmentEngine(store).score(UUID(spec_id), artifact.raw_content).overall_score
        except Exception:
            score = None

        try:
            saved = store.save_artifact(
                spec_id=UUID(spec_id),
                skill_id=skill_id,
                spec_name=artifact.spec_name,
                sections=artifact.sections,
                raw_content=artifact.raw_content,
                alignment_score=score,
            )
            artifact_history_id = saved["id"]
        except Exception:
            artifact_history_id = None

        if artifact_history_id:
            if artifact.renderer_type == "fabric":
                visual_url = f"{settings.base_url}/canvas?artifact_id={artifact_history_id}"
            elif artifact.renderer_type == "reveal":
                visual_url = f"{settings.base_url}/presentation/{artifact_history_id}"

        return {
            "skill_id": skill_id,
            "spec_name": artifact.spec_name,
            "spec_id": spec_id,
            "sections": artifact.sections,
            "raw_content": artifact.raw_content,
            "visual_url": visual_url,
            "artifact_id": artifact_history_id,
        }
    except Exception as e:
        raise HTTPException(500, str(e))


_SECTION_PROMPTS = {
    "summary": "Write a 2-3 sentence product summary for {name}. Positioning: {positioning}",
    "audience": "Describe the target audience for {name} in one focused paragraph. Use: {positioning}",
    "brand_personality": "Define the brand personality/voice for {name} in 2-3 sentences. Positioning: {positioning}",
    "tagline": "Write a punchy tagline (7 words or fewer) for {name}. Positioning: {positioning}",
    "differentiation": "List 2-3 key differentiators for {name} vs competitors. Positioning: {positioning}",
    "messages:headline": "Write 3 compelling benefit-led headlines for {name}. Positioning: {positioning}. Return as a bulleted list.",
    "messages:benefit": "Write 3 outcome-focused benefit statements for {name}. Positioning: {positioning}. Return as a bulleted list.",
    "messages:proof_point": "Write 3 credible proof points / stats for {name}. If no real stats exist in context, create plausible placeholders clearly marked [PLACEHOLDER]. Positioning: {positioning}. Return as a bulleted list.",
    "messages:qa_pair": "Write 3 common qa_pair handlers for {name}. Positioning: {positioning}. Return as a bulleted list.",
    "audiences": "Define 2 key buyer audiences for {name}. For each include: name, role/description, 3 pain points, 2 buying triggers, 2 qa_pairs. Positioning: {positioning}",
}


@app.post("/api/generate-section")
def generate_section(spec_id: str = Form(...), section: str = Form(...)):
    """Generate content for a specific missing section using the LLM."""
    try:
        h = store.get_spec(UUID(spec_id))
    except Exception:
        raise HTTPException(400, "Invalid spec_id")
    if not h:
        raise HTTPException(404, "Spec not found")

    template = _SECTION_PROMPTS.get(section)
    if not template:
        raise HTTPException(400, f"Unknown section: {section}")

    prompt = template.format(
        name=h.name,
        positioning=h.positioning or h.summary or "",
    )

    client = _get_oai_client()
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a B2B messaging strategist. Be specific, benefit-led, and concise."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.5,
        max_tokens=600,
    )
    store.record_token_usage(
        workspace_id="default",
        endpoint="generate-section",
        model="gpt-4o-mini",
        input_tokens=resp.usage.prompt_tokens,
        output_tokens=resp.usage.completion_tokens,
        cost_usd=estimate_cost_usd("gpt-4o-mini", resp.usage.prompt_tokens, resp.usage.completion_tokens),
    )
    return {"section": section, "content": resp.choices[0].message.content.strip()}


_SECTION_REGEN_PROMPTS = {
    "summary": "Rewrite the summary for {name}. Positioning: {positioning}. Keep it 2-3 sentences, compelling and specific.",
    "audience": "Rewrite the target audience description for {name}. Positioning: {positioning}. Be specific about role, company size, industry.",
    "tagline": "Create a new tagline for {name}. Positioning: {positioning}. Must be 7 words or fewer, memorable and ownable.",
    "differentiation": "Rewrite the differentiation for {name}. Positioning: {positioning}. List 2-3 specific ways this is better than alternatives.",
    "headline": "Write 3 new compelling headlines for {name}. Positioning: {positioning}. Benefit-led, specific, punchy. Return as bulleted list.",
    "subhead": "Write 3 new subheadlines for {name}. Positioning: {positioning}. Expand on headlines. Return as bulleted list.",
    "benefit": "Write 3 new benefit statements for {name}. Positioning: {positioning}. Outcome-focused with evidence. Return as bulleted list.",
    "proof_point": "Write 3 new proof points for {name}. Positioning: {positioning}. Quantified stats or customer evidence. Return as bulleted list.",
    "qa_pair": "Write 3 new qa_pair handlers for {name}. Positioning: {positioning}. Concise rebuttals. Return as bulleted list.",
    "social_proof": "Write 3 new social proof items for {name}. Positioning: {positioning}. Customer quotes, awards, G2 recognition. Return as bulleted list.",
}


@app.post("/api/generate-section-single")
def generate_section_single(spec_id: str = Query(...), section: str = Query(...)):
    """Regenerate a single section of an artifact using LLM."""
    try:
        h = store.get_spec(UUID(spec_id))
    except Exception:
        raise HTTPException(400, "Invalid spec_id")
    if not h:
        raise HTTPException(404, "Spec not found")

    messages = store.get_key_messages(h.id)
    messages_by_section = {}
    for m in messages:
        st = str(m.assertion_type)
        messages_by_section.setdefault(st, []).append(m.content)
    existing_msgs = "\n".join(f"- {st}: {msg}" for st, msgs in messages_by_section.items() for msg in msgs[:2])

    template = _SECTION_REGEN_PROMPTS.get(section.lower())
    if not template:
        template = _SECTION_PROMPTS.get(section, "Generate content for {name}. Positioning: {positioning}")

    prompt = template.format(
        name=h.name,
        positioning=h.positioning or h.summary or "",
    ) + f"\n\nExisting key messages:\n{existing_msgs}"

    client = _get_oai_client()
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a B2B messaging strategist. Be specific, benefit-led, and concise."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.6,
        max_tokens=600,
    )
    store.record_token_usage(
        workspace_id="default",
        endpoint="generate-section-single",
        model="gpt-4o-mini",
        input_tokens=resp.usage.prompt_tokens,
        output_tokens=resp.usage.completion_tokens,
        cost_usd=estimate_cost_usd("gpt-4o-mini", resp.usage.prompt_tokens, resp.usage.completion_tokens),
    )
    return {"section": section, "content": resp.choices[0].message.content.strip()}


# --- Spec Entry Improve / Variant Generation ---

@app.post("/api/entries/{entry_id}/improve")
@app.post("/api/messages/{msg_id}/improve")
def improve_message(msg_id: Optional[str] = None, entry_id: Optional[str] = None):
    """Suggest a stronger version of a assertion via LLM."""
    actual_id = entry_id or msg_id
    msg = _find_message(actual_id)
    if not msg:
        raise HTTPException(404, "Spec entry not found")
    spec = store.get_spec(msg.spec_id)
    positioning = spec.positioning if spec else ""

    client = _get_oai_client()
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a messaging and positioning expert. Rewrite the entry to be more specific, truth-grounded, and compelling. Return only the improved entry text — no preamble."},
            {"role": "user", "content": f"Section type: {msg.assertion_type}\nPositioning context: {positioning}\n\nOriginal entry:\n{msg.content}\n\nImproved version:"},
        ],
        temperature=0.6,
        max_tokens=300,
    )
    store.record_token_usage(
        workspace_id="default",
        endpoint="improve-message",
        model="gpt-4o-mini",
        input_tokens=resp.usage.prompt_tokens,
        output_tokens=resp.usage.completion_tokens,
        cost_usd=estimate_cost_usd("gpt-4o-mini", resp.usage.prompt_tokens, resp.usage.completion_tokens),
    )
    return {"original": msg.content, "improved": resp.choices[0].message.content.strip()}


@app.post("/api/entries/{entry_id}/generate-variant")
@app.post("/api/messages/{msg_id}/generate-variant")
def generate_variant(channel: str = Form(...), msg_id: Optional[str] = None, entry_id: Optional[str] = None):
    """Generate a channel-specific variant of a assertion."""
    actual_id = entry_id or msg_id
    msg = _find_message(actual_id)
    if not msg:
        raise HTTPException(404, "Spec entry not found")

    channel_guidance = {
        "linkedin": "LinkedIn post hook (under 150 chars, stops the scroll)",
        "email": "email subject line + opening hook (subject max 60 chars)",
        "twitter": "Twitter/X post (under 280 chars, punchy)",
        "paid": "paid ad headline + description (headline max 30 chars, description max 90 chars)",
        "landing": "landing page headline (benefit-led, max 10 words)",
        "blog": "blog post title and meta description",
    }
    guidance = channel_guidance.get(channel, f"{channel} version")

    client = _get_oai_client()
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": f"You are a B2B copywriter. Adapt the assertion for {channel}: {guidance}. Return only the adapted text."},
            {"role": "user", "content": f"Original entry:\n{msg.content}"},
        ],
        temperature=0.6,
        max_tokens=200,
    )
    variant_text = resp.choices[0].message.content.strip()
    store.record_token_usage(
        workspace_id="default",
        endpoint="generate-variant",
        model="gpt-4o-mini",
        input_tokens=resp.usage.prompt_tokens,
        output_tokens=resp.usage.completion_tokens,
        cost_usd=estimate_cost_usd("gpt-4o-mini", resp.usage.prompt_tokens, resp.usage.completion_tokens),
    )

    # Save the variant back to the entry
    variants = dict(msg.variants or {})
    variants[channel] = variant_text
    msg.variants = variants
    store.upsert_key_message(msg)

    return {"channel": channel, "variant": variant_text, "msg_id": actual_id, "entry_id": actual_id}


# --- Audience Generation ---

class GeneratePersonaRequest(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    spec_id: str = Field(validation_alias=AliasChoices('spec_id', 'canon_domain_id'), serialization_alias='spec_id')
    job_title: str



@app.post("/api/generate-audience")
def generate_audience(data: GeneratePersonaRequest):
    """Generate a full audience from a job title using LLM."""
    spec = store.get_spec(UUID(data.spec_id))
    if not spec:
        raise HTTPException(404, "Spec not found")

    client = _get_oai_client()
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a B2B buyer audience expert. Return JSON only."},
            {"role": "user", "content": f"""Generate a buyer audience for a '{data.job_title}' who might buy {spec.name}.

Context: {spec.positioning or spec.summary or 'B2B SaaS product'}

Return JSON:
{{
  "name": "<descriptive audience name like 'SMB CTO'>",
  "description": "<1-2 sentence role description>",
  "pain_points": ["<pain 1>", "<pain 2>", "<pain 3>"],
  "buying_triggers": ["<trigger 1>", "<trigger 2>"],
  "qa_pairs": ["<qa_pair 1>", "<qa_pair 2>"]
}}"""},
        ],
        response_format={"type": "json_object"},
        temperature=0.5,
        max_tokens=400,
    )
    store.record_token_usage(
        workspace_id="default",
        endpoint="generate-audience",
        model="gpt-4o-mini",
        input_tokens=resp.usage.prompt_tokens,
        output_tokens=resp.usage.completion_tokens,
        cost_usd=estimate_cost_usd("gpt-4o-mini", resp.usage.prompt_tokens, resp.usage.completion_tokens),
    )
    try:
        import json as _json
        audience_data = _json.loads(resp.choices[0].message.content)
    except Exception:
        raise HTTPException(500, "Failed to parse LLM audience response")

    audience = Audience(
        spec_id=spec.id,
        name=audience_data.get("name", data.job_title),
        description=audience_data.get("description", ""),
        pain_points=audience_data.get("pain_points", []),
        buying_triggers=audience_data.get("buying_triggers", []),
        qa_pairs=audience_data.get("qa_pairs", []),
    )
    store.upsert_audience(audience)
    return {
        "id": str(audience.id),
        "name": audience.name,
        "description": audience.description,
        "pain_points": audience.pain_points,
        "buying_triggers": audience.buying_triggers,
        "qa_pairs": audience.qa_pairs,
    }


# --- Tone Check ---

@app.post("/api/canon-domains/{domain_id}/check-tone")
@app.post("/api/specs/{spec_id}/check-tone")
def check_tone(spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    """Analyze assertions against brand_personality and flag mismatches."""
    actual_id = domain_id or spec_id
    spec = store.get_spec(UUID(actual_id))
    if not spec:
        raise HTTPException(404, "Spec domain not found")
    messages = store.get_key_messages(UUID(actual_id))
    if not messages:
        return {"warnings": [], "score": 100, "summary": "No entries to check"}
    if not spec.brand_personality:
        return {"warnings": [], "score": 100, "summary": "No brand personality defined — add one in Overview to enable tone checking"}

    samples = [m.content for m in messages[:12]]

    import json as _json
    client = _get_oai_client()
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a brand voice analyst. Return JSON only."},
            {"role": "user", "content": f"""Brand personality: {spec.brand_personality}

Spec entries:
{chr(10).join(f'- {m}' for m in samples)}

Identify which entries (if any) are inconsistent with the brand personality. Return JSON:
{{
  "score": <0-100 tone alignment score>,
  "summary": "<1 sentence overall assessment>",
  "warnings": [
    {{"message": "<entry text>", "issue": "<what's inconsistent>", "suggestion": "<how to fix>"}}
  ]
}}"""},
        ],
        response_format={"type": "json_object"},
        temperature=0.3,
        max_tokens=600,
    )
    store.record_token_usage(
        workspace_id="default",
        endpoint="check-tone",
        model="gpt-4o-mini",
        input_tokens=resp.usage.prompt_tokens,
        output_tokens=resp.usage.completion_tokens,
        cost_usd=estimate_cost_usd("gpt-4o-mini", resp.usage.prompt_tokens, resp.usage.completion_tokens),
    )
    try:
        result = _json.loads(resp.choices[0].message.content)
    except Exception:
        raise HTTPException(500, "Failed to parse tone check response")
    return result


# --- Snapshots ---

@app.post("/api/canon-domains/{domain_id}/snapshots")
@app.post("/api/specs/{spec_id}/snapshots")
def create_snapshot(data: dict = {}, spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    """Save a snapshot of the current framework state."""
    actual_id = domain_id or spec_id
    label = (data or {}).get("label", "")
    try:
        snap = store.create_snapshot(UUID(actual_id), label)
    except ValueError as e:
        raise HTTPException(404, str(e))
    return snap


@app.get("/api/canon-domains/{domain_id}/snapshots")
@app.get("/api/specs/{spec_id}/snapshots")
def list_snapshots(spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    actual_id = domain_id or spec_id
    return store.list_snapshots(UUID(actual_id))


@app.get("/api/snapshots/{snapshot_id}")
def get_snapshot(snapshot_id: str):
    snap = store.get_snapshot(UUID(snapshot_id))
    if not snap:
        raise HTTPException(404, "Snapshot not found")
    return snap


@app.get("/api/snapshots/{snapshot_id}/diff")
def get_snapshot_diff(snapshot_id: str):
    """Return a diff between a snapshot and the current framework state."""
    try:
        result = store.diff_snapshot(UUID(snapshot_id))
    except ValueError as e:
        raise HTTPException(404, str(e))
    return result


@app.delete("/api/snapshots/{snapshot_id}")
def delete_snapshot(snapshot_id: str):
    if not store.delete_snapshot(UUID(snapshot_id)):
        raise HTTPException(404, "Snapshot not found")
    return {"ok": True}


@app.post("/api/snapshots/{snapshot_id}/restore")
def restore_snapshot(snapshot_id: str):
    """Restore a framework to a snapshot state (replaces messages and audiences)."""
    snap = store.get_snapshot(UUID(snapshot_id))
    if not snap:
        raise HTTPException(404, "Snapshot not found")

    data = snap["snapshot_json"]
    spec_data = data.get("spec", {})
    spec_id = UUID(spec_data["id"])

    spec = store.get_spec(spec_id)
    if not spec:
        raise HTTPException(404, "Framework no longer exists")

    # Restore spec fields
    for field in ("name", "summary", "audience", "brand_personality", "positioning", "tagline", "differentiation"):
        if field in spec_data:
            setattr(spec, field, spec_data[field])
    store.upsert_spec(spec)

    # Replace messages
    for m in store.get_key_messages(spec_id):
        store.delete_key_message(m.id)
    for md in data.get("messages", []):
        try:
            msg = Assertion(
                spec_id=spec_id,
                assertion_type=AssertionType(md["assertion_type"].split(".")[-1] if "." in md["assertion_type"] else md["assertion_type"]),
                priority=md.get("priority", 3),
                content=md["content"],
                variants=md.get("variants", {}),
                audiences=md.get("audiences", []),
                channels=[Channel(c.split(".")[-1] if "." in c else c) for c in md.get("channels", ["all"])],
            )
            store.upsert_key_message(msg)
        except Exception:
            pass

    # Replace audiences
    for p in store.get_audiences(spec_id):
        store.delete_audience(p.id)
    for pd in data.get("audiences", []):
        audience = Audience(
            spec_id=spec_id,
            name=pd["name"],
            description=pd.get("description", ""),
            pain_points=pd.get("pain_points", []),
            buying_triggers=pd.get("buying_triggers", []),
            qa_pairs=pd.get("qa_pairs", []),
        )
        store.upsert_audience(audience)

    return {"ok": True, "spec_id": str(spec_id), "restored_from": snapshot_id}


# --- Artifact History ---

@app.post("/api/artifacts/save")
def save_artifact(data: dict):
    """Save a generated artifact to history."""
    try:
        spec_id = UUID(data["spec_id"])
    except Exception:
        raise HTTPException(400, "Invalid spec_id")
        
    try:
        from src.pipeline.alignment import AlignmentEngine
        score = AlignmentEngine(store).score(spec_id, data.get("raw_content", "")).overall_score
    except Exception:
        score = None

    record = store.save_artifact(
        spec_id=spec_id,
        skill_id=data.get("skill_id", ""),
        spec_name=data.get("spec_name", ""),
        sections=data.get("sections", {}),
        raw_content=data.get("raw_content", ""),
        alignment_score=score,
    )
    return record

@app.get("/api/artifacts/{artifact_id}/render")
def render_artifact(artifact_id: str, renderer: str = Query("fabric")):
    """Convert an artifact's sections to a specific renderer format."""
    from src.store import ArtifactHistoryModel
    import json
    aid = UUID(artifact_id)
    with store.session() as s:
        record = s.query(ArtifactHistoryModel).filter(ArtifactHistoryModel.id == str(aid)).first()
        if not record:
            raise HTTPException(404, "Artifact not found")
        sections = record.sections_json
        
        if renderer == "fabric":
            try:
                if "design_spec" in sections:
                    return json.loads(sections["design_spec"]) if isinstance(sections["design_spec"], str) else sections["design_spec"]
                
                positioning = (
                    sections.get("positioning")
                    or sections.get("positioning_statement")
                    or ""
                )
                assertions = (
                    sections.get("assertions")
                    or sections.get("key_messages_list")
                    or ""
                )
                tagline = sections.get("tagline", "")
                differentiation = sections.get("differentiation", "")
                audiences = sections.get("audiences", "")
                zones = [
                    {"type": "hero", "text": (tagline or record.spec_name)},
                    {"type": "positioning", "text": positioning},
                    {"type": "messages", "text": str(assertions)},
                ]
                if differentiation:
                    zones.append({"type": "differentiation", "text": str(differentiation)})
                if audiences:
                    zones.append({"type": "audiences", "text": str(audiences)})
                return {"zones": zones, "raw_sections": sections}
            except Exception as e:
                raise HTTPException(500, f"Failed to format for fabric: {str(e)}")
        
        raise HTTPException(400, "Unsupported renderer")

class DesignSpecUpdate(BaseModel):
    design_spec: dict

@app.get("/api/artifacts/{artifact_id}/design_spec")
def get_artifact_design_spec(artifact_id: str):
    """Fetch the Fabric.js design spec for an artifact, falling back to a generated template spec."""
    from src.store import ArtifactHistoryModel
    try:
        aid = UUID(artifact_id)
    except Exception:
        raise HTTPException(400, "Invalid artifact ID")
    
    with store.session() as s:
        record = s.query(ArtifactHistoryModel).filter(ArtifactHistoryModel.id == str(aid)).first()
        if not record:
            raise HTTPException(404, "Artifact not found")
        
        sections = record.sections_json or {}
        if "design_spec" in sections and sections["design_spec"]:
            spec = sections["design_spec"]
            return json.loads(spec) if isinstance(spec, str) else spec
        
        # No design_spec saved yet, generate from template
        from src.pipeline.generator import ArtifactGenerator
        generator = ArtifactGenerator(store, skills)
        
        # Skills mapping
        artifact_type = record.skill_id
        template = generator._get_template(artifact_type)
        if template:
            spec = store.get_spec(UUID(record.spec_id)) if record.spec_id else None
            if spec:
                messages = store.get_key_messages(spec.id)
                audiences = store.get_audiences(spec.id)
                visual_context = generator._build_visual_context(
                    record.spec_id, template, artifact_type, spec, messages, audiences
                )
                spec = generator._fallback_design_spec(template, generator._build_context(spec, messages, audiences, {}), visual_context)
                return spec
        
        # Create a minimal fallback design spec if no template matches
        return {
            "version": "2.0",
            "artifact_type": record.skill_id,
            "template_id": record.skill_id,
            "zones": [],
            "canvas_width": 850,
            "canvas_height": 1100,
            "background": {"type": "solid", "color": "#FFFFFF"},
        }

@app.patch("/api/artifacts/{artifact_id}/design_spec")
def update_artifact_design_spec(artifact_id: str, data: DesignSpecUpdate):
    """Save the updated Fabric.js design spec back to the artifact."""
    from src.store import ArtifactHistoryModel
    aid = UUID(artifact_id)
    with store.session() as s:
        record = s.query(ArtifactHistoryModel).filter(ArtifactHistoryModel.id == str(aid)).first()
        if not record:
            raise HTTPException(404, "Artifact not found")
        
        # We store the updated spec back into the 'design_spec' section
        sections = dict(record.sections_json)
        sections["design_spec"] = data.design_spec
        record.sections_json = sections
        s.commit()

    return {"ok": True}


# --- Brand Settings ---

@app.get("/api/workspaces/{workspace_id}/brand")
def get_workspace_brand(workspace_id: str):
    """Get brand settings for a workspace."""
    settings = store.get_brand_settings(workspace_id)
    if not settings:
        # Return defaults if not configured
        return {
            "workspace_id": workspace_id,
            "primary_color": "#1e293b",
            "secondary_color": "#3b82f6",
            "accent_color": "#f59e0b",
            "background_color": "#ffffff",
            "text_color": "#1e293b",
            "font_heading": "Inter",
            "font_body": "Inter",
            "logo_path": None,
            "custom_fonts": [],
        }
    return settings


@app.patch("/api/workspaces/{workspace_id}/brand")
def update_workspace_brand(workspace_id: str, data: dict):
    """Update brand settings for a workspace."""
    result = store.upsert_brand_settings(workspace_id, **data)
    return result


# --- Design Spec Reset ---

@app.post("/api/artifacts/{artifact_id}/design_spec/reset")
def reset_artifact_design_spec(artifact_id: str):
    """Reset design spec to AI-generated version."""
    try:
        aid = UUID(artifact_id)
    except Exception:
        raise HTTPException(400, "Invalid artifact ID")

    from src.store import ArtifactHistoryModel
    with store.session() as s:
        record = s.query(ArtifactHistoryModel).filter(ArtifactHistoryModel.id == str(aid)).first()
        if not record:
            raise HTTPException(404, "Artifact not found")

        # Remove the design_spec to force regeneration
        sections = dict(record.sections_json)
        if "design_spec" in sections:
            del sections["design_spec"]
            record.sections_json = sections
            s.commit()

        # Regenerate from AI
        try:
            from src.pipeline.generator import ArtifactGenerator
            generator = ArtifactGenerator(store=store, openai_client=_get_oai_client())
            # Get spec for this artifact
            spec = store.get_spec(UUID(record.spec_id)) if record.spec_id else None
            if spec:
                _, updated_sections = generator.generate_artifact(
                    spec_id=record.spec_id,
                    skill_id=record.skill_id,
                    custom_context=record.custom_context,
                    renderer="fabric"
                )
                sections = dict(record.sections_json)
                if "design_spec" in updated_sections:
                    sections["design_spec"] = updated_sections["design_spec"]
                    record.sections_json = sections
                    s.commit()
        except Exception as e:
            log.warning("Failed to regenerate design spec: %s", e)

    return {"ok": True, "reset": True}


@app.get("/api/recent-artifacts")
def list_recent_artifacts(limit: int = Query(5, le=20)):
    return store.list_recent_artifacts(limit=limit)


@app.get("/api/canon-domains/{domain_id}/artifacts")
@app.get("/api/specs/{spec_id}/artifacts")
def list_spec_artifacts(spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    actual_id = domain_id or spec_id
    return store.list_artifacts(UUID(actual_id))


@app.get("/api/artifacts/{artifact_id}")
def get_artifact(artifact_id: str):
    record = store.get_artifact(UUID(artifact_id))
    if not record:
        raise HTTPException(404, "Artifact not found")
    return record


# --- Artifact Ratings ---

class ArtifactRateRequest(BaseModel):
    rating: int  # 1-5
    tag: str = ""  # "good" or "bad", auto-inferred if empty
    rated_by: str = ""
    notes: str = ""


@app.post("/api/artifacts/{artifact_id}/rate")
def rate_artifact(artifact_id: str, data: ArtifactRateRequest):
    """Rate an artifact (1-5 stars or good/bad). Updates chunk usage stats."""
    try:
        aid = UUID(artifact_id)
    except Exception:
        raise HTTPException(400, "Invalid artifact ID")
    result = store.rate_artifact(
        artifact_id=str(aid),
        rating=data.rating,
        tag=data.tag,
        rated_by=data.rated_by,
        notes=data.notes,
    )
    return result


@app.get("/api/artifacts/{artifact_id}/ratings")
def get_artifact_ratings(artifact_id: str):
    """Get all ratings for an artifact."""
    try:
        aid = UUID(artifact_id)
    except Exception:
        raise HTTPException(400, "Invalid artifact ID")
    return {"artifact_id": artifact_id, "ratings": store.get_artifact_rating(str(aid))}


# --- Usage Heatmap & Coverage ---

@app.get("/api/canon-domains/{domain_id}/heatmap")
@app.get("/api/specs/{spec_id}/heatmap")
def get_heatmap(spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    """Get chunk usage heatmap for a spec."""
    actual_id = domain_id or spec_id
    try:
        hid = UUID(actual_id)
    except Exception:
        raise HTTPException(400, "Invalid ID")
    return store.get_chunk_usage_heatmap(hid)


@app.get("/api/canon-domains/{domain_id}/coverage")
@app.get("/api/specs/{spec_id}/coverage")
def get_coverage(spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    """Get spec coverage report."""
    actual_id = domain_id or spec_id
    try:
        hid = UUID(actual_id)
    except Exception:
        raise HTTPException(400, "Invalid ID")
    return store.get_spec_coverage(hid)


@app.patch("/api/artifacts/{artifact_id}")
def update_artifact(artifact_id: str, data: dict):
    """Update sections or raw_content of an existing artifact."""
    record = store.get_artifact(UUID(artifact_id))
    if not record:
        raise HTTPException(404, "Artifact not found")

    sections = record.get("sections", {})
    if "sections" in data and isinstance(data["sections"], dict):
        sections.update(data["sections"])

    from src.store import ArtifactHistoryModel
    from uuid import UUID as UUID_
    with store.session() as s:
        row = s.get(ArtifactHistoryModel, artifact_id)
        if row:
            row.sections_json = sections
            if "raw_content" in data:
                row.raw_content = data["raw_content"]
            s.commit()

    return {"ok": True, "updated_id": artifact_id}


class ArtifactStatusUpdate(BaseModel):
    status: str
    notes: Optional[str] = None


_ARTIFACT_STATUS_TRANSITIONS = {
    "draft": {"internal_review"},
    "internal_review": {"draft", "approved"},
    "approved": {"internal_review"},
}


@app.patch("/api/artifacts/{artifact_id}/status")
def update_artifact_status(artifact_id: str, data: ArtifactStatusUpdate):
    """Transition an artifact through its review lifecycle.

    Valid transitions:
      draft → internal_review
      internal_review → draft | approved
      approved → internal_review
    """
    try:
        valid_statuses = {s.value for s in ArtifactStatus}
    except Exception:
        valid_statuses = {"draft", "internal_review", "approved"}

    if data.status not in valid_statuses:
        raise HTTPException(400, f"Invalid status '{data.status}'. Must be one of: {sorted(valid_statuses)}")

    record = store.get_artifact(UUID(artifact_id))
    if not record:
        raise HTTPException(404, "Artifact not found")

    current = record.get("status", "draft")
    allowed = _ARTIFACT_STATUS_TRANSITIONS.get(current, set())
    if data.status not in allowed:
        raise HTTPException(409, f"Cannot transition artifact from '{current}' to '{data.status}'. Allowed: {sorted(allowed)}")

    if not store.update_artifact_status(UUID(artifact_id), data.status):
        raise HTTPException(500, "Failed to update artifact status")

    return {"ok": True, "id": artifact_id, "status": data.status}


# --- DOCX Download ---

@app.get("/api/artifacts/{artifact_id}/docx")
def download_artifact_docx(artifact_id: str):
    """Download a saved artifact as a DOCX file."""
    record = store.get_artifact(UUID(artifact_id))
    if not record:
        raise HTTPException(404, "Artifact not found")

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from io import BytesIO
        from fastapi.responses import StreamingResponse
    except ImportError:
        raise HTTPException(500, "python-docx not installed")

    doc = DocxDocument()
    doc.add_heading(f"{record['spec_name']} — {record['skill_id'].replace('_', ' ').title()}", 0)

    for section_key, section_value in record["sections"].items():
        doc.add_heading(section_key.replace("_", " ").title(), 2)
        doc.add_paragraph(str(section_value))

    if record.get("raw_content"):
        doc.add_heading("Full Content", 2)
        doc.add_paragraph(record["raw_content"])

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{record['spec_name'].replace(' ', '_')}_{record['skill_id']}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/artifacts/{artifact_id}/pdf", response_class=HTMLResponse)
def download_artifact_pdf(artifact_id: str):
    """Return a print-ready HTML page that auto-triggers browser print dialog."""
    record = store.get_artifact(UUID(artifact_id))
    if not record:
        raise HTTPException(404, "Artifact not found")

    sections_html = ""
    for key, value in record["sections"].items():
        sections_html += f"""
        <div class="section">
            <h2>{key.replace('_', ' ').title()}</h2>
            <div class="section-body">{str(value).replace(chr(10), '<br>')}</div>
        </div>"""

    if record.get("raw_content"):
        sections_html += f"""
        <div class="section">
            <h2>Full Content</h2>
            <div class="section-body">{record['raw_content'].replace(chr(10), '<br>')}</div>
        </div>"""

    html = f"""<!doctype html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{record['spec_name']} — {record['skill_id'].replace('_', ' ').title()}</title>
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
        * {{ box-sizing: border-box; margin: 0; padding: 0; }}
        body {{ font-family: 'Inter', sans-serif; color: #0f172a; padding: 40px; max-width: 800px; margin: 0 auto; }}
        h1 {{ font-size: 28px; font-weight: 700; margin-bottom: 4px; }}
        .subtitle {{ color: #64748b; font-size: 14px; margin-bottom: 32px; padding-bottom: 16px; border-bottom: 2px solid #e2e8f0; }}
        .section {{ margin-bottom: 28px; page-break-inside: avoid; }}
        h2 {{ font-size: 13px; font-weight: 700; text-transform: uppercase; letter-spacing: 0.08em; color: #6366f1; margin-bottom: 10px; }}
        .section-body {{ font-size: 15px; line-height: 1.7; color: #334155; }}
        @media print {{ body {{ padding: 0; }} }}
        </style>
    </head>
    <body>
        <h1>{record['spec_name']}</h1>
        <div class="subtitle">{record['skill_id'].replace('_', ' ').title()} · Generated {record['created_at'][:10]}</div>
        {sections_html}
        <script>window.onload = () => window.print();</script>
    </body>
    </html>"""
    return HTMLResponse(content=html)


# --- Completeness in specs list ---

def _completeness_score(spec: Spec) -> int:
    """Return 0-100 completeness score for a spec."""
    messages = store.get_key_messages(spec.id)
    audiences = store.get_audiences(spec.id)
    score = 0
    if spec.name: score += 10
    if spec.summary: score += 10
    if spec.audience: score += 10
    if spec.brand_personality: score += 10
    if spec.positioning: score += 15
    if spec.tagline: score += 10
    if spec.differentiation: score += 10
    headlines = [m for m in messages if str(m.assertion_type).endswith("headline")]
    if headlines: score += 10
    benefits = [m for m in messages if str(m.assertion_type).endswith("benefit")]
    if benefits: score += 10
    if audiences: score += 5
    return min(score, 100)


# --- Health + Metrics ---

@app.get("/health")
def health():
    """Production health check — returns 200 when the DB is reachable."""
    try:
        store.list_workspaces()
        db_ok = True
    except Exception:
        db_ok = False
    return {
        "status": "ok" if db_ok else "degraded",
        "db": "ok" if db_ok else "error",
        "uptime_seconds": round(time.time() - _start_time),
        "version": "0.5.0",
        "auth_enabled": settings.auth_enabled,
    }


@app.get("/api/metrics")
def get_metrics(auth: AuthContext = Depends(require_read)):
    """Request metrics + token usage summary."""
    top = sorted(_metrics.items(), key=lambda x: x[1]["requests"], reverse=True)[:20]
    return {
        "uptime_seconds": round(time.time() - _start_time),
        "endpoints": [
            {
                "path": k,
                "requests": v["requests"],
                "errors": v["errors"],
                "avg_latency_ms": round(v["total_ms"] / max(v["requests"], 1), 1),
            }
            for k, v in top
        ],
        "token_usage": store.get_token_usage_summary(
            workspace_id=auth.workspace_id if settings.auth_enabled else None
        ),
    }


# --- Workspaces ---

class WorkspaceCreate(BaseModel):
    slug: str
    name: str
    max_token_budget: int = 0


@app.get("/api/workspaces")
def list_workspaces(auth: AuthContext = Depends(require_read)):
    workspaces = store.list_workspaces()
    for ws in workspaces:
        ws["token_usage"] = store.get_token_usage_summary(workspace_id=ws["id"])
        specs = store.list_specs(workspace_id=ws["id"])
        ws["spec_count"] = len(specs)
    return workspaces


@app.post("/api/workspaces")
def create_workspace(data: WorkspaceCreate, auth: AuthContext = Depends(require_write)):
    if not data.slug.replace("-", "").replace("_", "").isalnum():
        raise HTTPException(400, "Slug must be alphanumeric with hyphens/underscores only")
    existing = store.get_workspace(data.slug)
    if existing:
        raise HTTPException(409, f"Workspace slug '{data.slug}' already exists")
    return store.create_workspace(data.slug, data.name, data.max_token_budget)


@app.get("/api/workspaces/{workspace_id}")
def get_workspace(workspace_id: str, auth: AuthContext = Depends(require_read)):
    ws = store.get_workspace(workspace_id)
    if not ws:
        raise HTTPException(404, "Workspace not found")
    ws["token_usage"] = store.get_token_usage_summary(workspace_id=ws["id"])
    ws["spec_count"] = len(store.list_specs(workspace_id=ws["id"]))
    return ws


class WorkspaceUpdate(BaseModel):
    name: Optional[str] = None
    max_token_budget: Optional[int] = None


@app.patch("/api/workspaces/{workspace_id}")
def update_workspace(workspace_id: str, data: WorkspaceUpdate, auth: AuthContext = Depends(require_write)):
    updates = data.model_dump(exclude_none=True)
    result = store.update_workspace(workspace_id, **updates)
    if not result:
        raise HTTPException(404, "Workspace not found")
    return result


# --- Brand Settings & logo ---

class BrandSettingsUpdate(BaseModel):
    primary_color: Optional[str] = None
    secondary_color: Optional[str] = None
    accent_color: Optional[str] = None
    background_color: Optional[str] = None
    text_color: Optional[str] = None
    font_heading: Optional[str] = None
    font_body: Optional[str] = None


@app.get("/api/workspaces/{workspace_id}/brand")
def get_brand_settings(workspace_id: str):
    """Fetch brand settings for a workspace, returning defaults if not found."""
    brand = store.get_brand_settings(workspace_id)
    if not brand:
        from src.models import BrandSettings
        brand = BrandSettings(workspace_id=workspace_id)
    return brand.model_dump()


@app.post("/api/workspaces/{workspace_id}/brand")
def update_brand_settings(workspace_id: str, data: BrandSettingsUpdate):
    """Upsert brand settings for a workspace."""
    updates = data.model_dump(exclude_none=True)
    brand = store.upsert_brand_settings(workspace_id, **updates)
    return brand.model_dump()


@app.post("/api/workspaces/{workspace_id}/brand/logo")
def upload_brand_logo(workspace_id: str, file: UploadFile = File(...)):
    """Upload a PNG or SVG logo for a workspace."""
    ext = file.filename.split(".")[-1].lower() if file.filename else "png"
    if ext not in ("png", "svg", "jpg", "jpeg"):
        raise HTTPException(400, "Only PNG, SVG, and JPG logos are supported.")
    
    brand_dir = DATA_DIR / "brand" / workspace_id
    brand_dir.mkdir(parents=True, exist_ok=True)
    
    file_path = brand_dir / f"logo.{ext}"
    try:
        with open(file_path, "wb") as f:
            f.write(file.file.read())
    except Exception as e:
        raise HTTPException(500, f"Failed to save logo file: {e}")
        
    rel_path = f"/api/workspaces/{workspace_id}/brand/logo"
    store.upsert_brand_settings(workspace_id, logo_path=rel_path)
    return {"logo_path": rel_path}


@app.get("/api/workspaces/{workspace_id}/brand/logo")
def get_brand_logo(workspace_id: str):
    """Serve the brand logo file."""
    brand_dir = DATA_DIR / "brand" / workspace_id
    files = list(brand_dir.glob("logo.*"))
    if not files:
        raise HTTPException(404, "Logo not found")
    return FileResponse(files[0])


# --- Template Registry ---

@app.get("/api/templates")
def list_templates():
    """List registered templates."""
    from src.design.template_registry import TemplateRegistry
    registry = TemplateRegistry()
    return registry.list_templates()


@app.get("/api/templates/{artifact_type}")
def get_template(artifact_type: str):
    """Get detailed zone layout template."""
    from src.design.template_registry import TemplateRegistry
    registry = TemplateRegistry()
    template = registry.get_template(artifact_type)
    if not template:
        raise HTTPException(404, "Template not found")
    return template.model_dump()


# --- API Key Management ---

class ApiKeyCreate(BaseModel):
    name: str
    workspace_id: str = "default"
    scopes: list[str] = ["read", "write"]


@app.post("/api/api-keys")
def create_api_key(data: ApiKeyCreate, auth: AuthContext = Depends(require_write)):
    valid_scopes = {"read", "write", "admin"}
    bad = {s for s in data.scopes if s not in valid_scopes and not s.startswith("dept:")}
    if bad:
        raise HTTPException(400, f"Invalid scopes: {bad}. Use: {valid_scopes} or dept:<name>")
    plaintext, key_hash = generate_api_key()
    record = store.create_api_key(
        key_hash=key_hash,
        name=data.name,
        workspace_id=data.workspace_id,
        scopes=data.scopes,
    )
    # Return plaintext key ONCE — it won't be shown again
    record["key"] = plaintext
    record["warning"] = "Save this key now — it will not be shown again."
    return record


@app.get("/api/api-keys")
def list_api_keys(workspace_id: Optional[str] = None, auth: AuthContext = Depends(require_read)):
    wid = workspace_id or (auth.workspace_id if settings.auth_enabled else None)
    return store.list_api_keys(workspace_id=wid)


@app.delete("/api/api-keys/{key_id}")
def revoke_api_key(key_id: str, auth: AuthContext = Depends(require_write)):
    if not store.revoke_api_key(key_id):
        raise HTTPException(404, "API key not found")
    return {"ok": True, "revoked": key_id}


# --- Token Usage ---

@app.get("/api/token-usage")
def get_token_usage(workspace_id: Optional[str] = None, auth: AuthContext = Depends(require_read)):
    wid = workspace_id or (auth.workspace_id if settings.auth_enabled else None)
    return store.get_token_usage_summary(workspace_id=wid)


@app.get("/api/cost-estimate")
def cost_estimate(model: str = "gpt-4o-mini", input_tokens: int = 0, output_tokens: int = 0):
    """Return a cost estimate for a given token count before running LLM."""
    cost = estimate_cost_usd(model, input_tokens, output_tokens)
    rates = settings.pricing.get(model, {})
    return {
        "model": model,
        "input_tokens": input_tokens,
        "output_tokens": output_tokens,
        "estimated_cost_usd": round(cost, 6),
        "rates_per_1m": rates,
    }


_ARTIFACT_SECTION_META = {
    "headline":     {"label": "Headline",     "color": "#6366f1", "bg": "rgba(99,102,241,.12)"},
    "subhead":      {"label": "Subhead",      "color": "#8b5cf6", "bg": "rgba(139,92,246,.12)"},
    "benefit":      {"label": "Benefit",      "color": "#10b981", "bg": "rgba(16,185,129,.12)"},
    "use_case":     {"label": "Use Case",     "color": "#06b6d4", "bg": "rgba(6,182,212,.12)"},
    "proof_point":  {"label": "Proof Point",  "color": "#3b82f6", "bg": "rgba(59,130,246,.12)"},
    "qa_pair":    {"label": "QAPair",    "color": "#ef4444", "bg": "rgba(239,68,68,.12)"},
    "social_proof": {"label": "Social Proof", "color": "#f59e0b", "bg": "rgba(245,158,11,.12)"},
    "positioning":  {"label": "Positioning",  "color": "#64748b", "bg": "rgba(100,116,139,.12)"},
}
_ARTIFACT_SECTION_ORDER = ["headline", "subhead", "benefit", "use_case", "proof_point", "qa_pair", "social_proof", "positioning"]
_ARTIFACT_TYPE_LABELS = {
    "one_pager": "One Pager",
    "social_posts": "Social Posts",
    "email_template": "Email Template",
    "battlecard": "Battlecard",
    "email_sequence": "Email Sequence",
}


@app.get("/presentation/{artifact_id}", response_class=HTMLResponse)
def serve_presentation(artifact_id: str):
    """Serve a generated reveal.js presentation."""
    try:
        from src.store import ArtifactHistoryModel
        from src.rendering.renderer import get_renderer
        aid = UUID(artifact_id)
        with store.session() as s:
            record = s.query(ArtifactHistoryModel).filter(ArtifactHistoryModel.id == str(aid)).first()
            if not record:
                raise HTTPException(404, "Artifact not found")
            
            renderer = get_renderer("reveal")
            context = {"spec_name": record.spec_name or "Untitled"}
            
            # Try to load workspace brand if possible (requires joining to spec)
            from src.models import Spec
            spec = s.query(Spec).filter(Spec.id == record.spec_id).first()
            if spec and spec.workspace_id:
                brand = store.get_workspace_brand(spec.workspace_id)
                if brand:
                    context["brand_settings"] = brand
                    
            output = renderer.render_reveal(record.sections_json, context)
            return HTMLResponse(content=output.content)
    except Exception as e:
        raise HTTPException(500, f"Error rendering presentation: {e}")


@app.get("/artifact/{artifact_type}/{spec_id}", response_class=HTMLResponse)
def serve_artifact(
    request: Request,
    artifact_type: str,
    spec_id: str,
    stage: str = "awareness",
    channels: str = "linkedin",
    competitor: str = "",
    auth: AuthContext = Depends(get_auth_context),
):
    """Serve a standalone HTML artifact page for a message spec."""
    if settings.auth_enabled and "read" not in auth.scopes:
        raise HTTPException(403, "Read scope required to view artifacts.")
    try:
        hid = UUID(spec_id)
    except ValueError:
        raise HTTPException(400, "Invalid spec_id UUID")

    valid_types = list(_ARTIFACT_TYPE_LABELS.keys())
    if artifact_type not in valid_types:
        raise HTTPException(400, f"Unknown artifact_type. Use: {', '.join(valid_types)}")

    spec = store.get_spec(hid)
    if not spec:
        raise HTTPException(404, "Spec not found")

    messages = store.get_key_messages(hid)
    audiences = store.get_audiences(hid)

    if artifact_type == "one_pager":
        grouped: dict[str, list] = {}
        for m in messages:
            st = str(m.assertion_type).split(".")[-1].lower().replace(" ", "_")
            grouped.setdefault(st, []).append(m.content)
        synced = spec.last_synced.strftime("%Y-%m-%d") if spec.last_synced else "—"
        return templates.TemplateResponse(request, "artifact_visual.html", {
            "spec": spec,
            "domain": spec,
            "grouped": grouped,
            "audiences": audiences,
            "section_meta": _ARTIFACT_SECTION_META,
            "section_order": _ARTIFACT_SECTION_ORDER,
            "message_count": len(messages),
            "entry_count": len(messages),
            "audience_count": len(audiences),
            "artifact_type_label": "One Pager",
            "synced_date": synced,
        })

    grouped_legacy: dict[str, list] = {}
    for m in messages:
        key = str(m.assertion_type).replace("_", " ").title()
        grouped_legacy.setdefault(key, []).append(m.content)

    if artifact_type == "social_posts":
        target = channels.split(",")
        html = _render_social_posts(spec, messages, target)
    elif artifact_type == "battlecard":
        html = _render_battlecard(spec, messages, competitor or "Competitor")
    elif artifact_type == "email_sequence":
        html = _render_email_sequence(spec, messages)
    else:
        html = _render_email_template(spec, messages, stage)

    return HTMLResponse(content=html)


_SECTION_META = {
    "Headline":    {"icon": "✦", "color": "#6366f1", "bg": "#eef2ff"},
    "Subhead":     {"icon": "◈", "color": "#8b5cf6", "bg": "#f5f3ff"},
    "Benefit":     {"icon": "◉", "color": "#059669", "bg": "#ecfdf5"},
    "Use Case":    {"icon": "⬡", "color": "#0891b2", "bg": "#ecfeff"},
    "Proof Point": {"icon": "◆", "color": "#0284c7", "bg": "#e0f2fe"},
    "QAPair":   {"icon": "◇", "color": "#dc2626", "bg": "#fef2f2"},
    "Social Proof":{"icon": "★", "color": "#d97706", "bg": "#fffbeb"},
    "Positioning": {"icon": "◎", "color": "#475569", "bg": "#f8fafc"},
}

_BASE_STYLES = """
  @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
  *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
  :root { --bg: #f1f5f9; --fg: #0f172a; --card-bg: #fff; --card-border: #e2e8f0; --muted: #64748b; --muted-2: #94a3b8; }
  body.dark { --bg: #0f172a; --fg: #f1f5f9; --card-bg: #1e293b; --card-border: #334155; --muted: #94a3b8; --muted-2: #64748b; }
  body { font-family: 'Inter', system-ui, sans-serif; background: var(--bg); color: var(--fg); -webkit-font-smoothing: antialiased; transition: background 0.2s, color 0.2s; }
  .theme-toggle { position: fixed; top: 16px; right: 16px; background: var(--card-bg); border: 1px solid var(--card-border); border-radius: 8px; padding: 6px 12px; font-size: 12px; cursor: pointer; color: var(--muted); z-index: 50; }
  @media print { .theme-toggle { display: none; } .page { padding: 0; } .hero { background: #0f172a !important; -webkit-print-color-adjust: exact; print-color-adjust: exact; } .card { box-shadow: none; border: 1px solid #ddd; break-inside: avoid; } }

  .page { max-width: 900px; margin: 0 auto; padding: 40px 24px 80px; }
  .hero { background: linear-gradient(135deg, #0f172a 0%, #1e1b4b 50%, #0f172a 100%); border-radius: 20px; padding: 52px 48px; margin-bottom: 24px; position: relative; overflow: hidden; }
  .hero::before { content: ''; position: absolute; inset: 0; background: radial-gradient(ellipse at 70% 50%, rgba(99,102,241,0.25) 0%, transparent 60%); }
  .hero-label { display: inline-flex; align-items: center; gap: 6px; background: rgba(255,255,255,0.1); border: 1px solid rgba(255,255,255,0.15); color: rgba(255,255,255,0.7); font-size: 11px; font-weight: 600; letter-spacing: 0.1em; text-transform: uppercase; padding: 5px 12px; border-radius: 100px; margin-bottom: 20px; }
  .hero h1 { font-size: 42px; font-weight: 800; color: #fff; line-height: 1.1; letter-spacing: -0.02em; margin-bottom: 12px; position: relative; }
  .hero-tagline { font-size: 18px; color: rgba(255,255,255,0.65); font-weight: 400; line-height: 1.5; margin-bottom: 24px; position: relative; max-width: 580px; }
  .hero-audience { display: inline-block; background: rgba(99,102,241,0.25); border: 1px solid rgba(99,102,241,0.4); color: #a5b4fc; font-size: 13px; font-weight: 500; padding: 6px 14px; border-radius: 100px; position: relative; }
  .card { background: #fff; border: 1px solid #e2e8f0; border-radius: 16px; padding: 32px; margin-bottom: 20px; box-shadow: 0 1px 3px rgba(0,0,0,0.06), 0 4px 16px rgba(0,0,0,0.04); }
  .card-label { font-size: 11px; font-weight: 700; letter-spacing: 0.08em; text-transform: uppercase; color: #94a3b8; margin-bottom: 16px; display: flex; align-items: center; gap: 8px; }
  .card-label::after { content: ''; flex: 1; height: 1px; background: #f1f5f9; }
  .positioning-text { font-size: 17px; color: #334155; line-height: 1.7; font-weight: 400; }
  .diff-text { margin-top: 16px; font-size: 14px; color: #64748b; line-height: 1.6; padding-left: 16px; border-left: 3px solid #e2e8f0; }
  .messages-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 20px; }
  @media (max-width: 640px) { .messages-grid { grid-template-columns: 1fr; } }
  .section-block { border-radius: 12px; padding: 20px; }
  .section-icon { font-size: 14px; margin-right: 6px; }
  .section-title { font-size: 11px; font-weight: 700; letter-spacing: 0.08em; text-transform: uppercase; margin-bottom: 12px; display: flex; align-items: center; }
  .msg-item { display: flex; gap: 10px; align-items: flex-start; padding: 8px 0; border-bottom: 1px solid rgba(0,0,0,0.05); }
  .msg-item:last-child { border-bottom: none; padding-bottom: 0; }
  .msg-num { min-width: 20px; height: 20px; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-size: 10px; font-weight: 700; margin-top: 1px; flex-shrink: 0; }
  .msg-text { font-size: 13.5px; color: #334155; line-height: 1.5; }
  .audiences-grid { display: grid; grid-template-columns: repeat(auto-fill, minmax(220px, 1fr)); gap: 16px; }
  .audience-card { border: 1px solid #e2e8f0; border-radius: 12px; padding: 20px; background: #fafafa; }
  .audience-avatar { width: 40px; height: 40px; border-radius: 12px; background: linear-gradient(135deg, #6366f1, #8b5cf6); display: flex; align-items: center; justify-content: center; color: #fff; font-weight: 700; font-size: 15px; margin-bottom: 12px; }
  .audience-name { font-size: 15px; font-weight: 600; color: #0f172a; margin-bottom: 4px; }
  .audience-desc { font-size: 12.5px; color: #64748b; line-height: 1.5; margin-bottom: 12px; }
  .pain-tag { display: inline-block; background: #fef2f2; color: #dc2626; font-size: 11px; font-weight: 500; padding: 3px 8px; border-radius: 6px; margin: 2px 2px 0 0; }
  .footer { text-align: center; margin-top: 32px; }
  .footer-badge { display: inline-flex; align-items: center; gap: 6px; background: #fff; border: 1px solid #e2e8f0; color: #94a3b8; font-size: 11px; font-weight: 500; padding: 6px 14px; border-radius: 100px; }
  .post-card { border: 1px solid #e2e8f0; border-radius: 14px; padding: 24px; margin-bottom: 16px; background: #fff; box-shadow: 0 1px 3px rgba(0,0,0,0.05); }
  .post-badge { display: inline-flex; align-items: center; gap-4px; font-size: 11px; font-weight: 600; padding: 4px 10px; border-radius: 100px; margin-right: 6px; margin-bottom: 14px; }
  .post-text { font-size: 15px; color: #1e293b; line-height: 1.7; white-space: pre-wrap; }
  .email-field { padding: 20px 0; border-bottom: 1px solid #f1f5f9; }
  .email-field:last-child { border-bottom: none; }
  .email-label { font-size: 10px; font-weight: 700; letter-spacing: 0.1em; text-transform: uppercase; color: #94a3b8; margin-bottom: 8px; }
  .email-value { font-size: 15px; color: #1e293b; line-height: 1.6; }
  .email-value.subject { font-size: 18px; font-weight: 600; }
  .email-value.cta { font-size: 16px; font-weight: 600; color: #6366f1; }
  .stage-tabs { display: flex; gap: 8px; margin-bottom: 24px; flex-wrap: wrap; }
  .stage-tab { padding: 6px 16px; border-radius: 100px; font-size: 12px; font-weight: 600; border: 1.5px solid #e2e8f0; color: #64748b; text-decoration: none; }
  .stage-tab.active { background: #0f172a; color: #fff; border-color: #0f172a; }
"""


def _base_html(title: str, body: str, extra_styles: str = "") -> str:
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{title}</title>
  <style>{_BASE_STYLES}{extra_styles}</style>
</head>
<body>
  <button class="theme-toggle" onclick="document.body.classList.toggle('dark')">&#9680; Toggle theme</button>
  <div class="page">
{body}
  </div>
  <script>
    if (window.matchMedia && window.matchMedia('(prefers-color-scheme: dark)').matches) {{
      document.body.classList.add('dark');
    }}
  </script>
</body>
</html>"""


def _render_one_pager(spec, grouped: dict, audiences: list) -> str:
    last_synced = spec.last_synced.strftime("%Y-%m-%d") if spec.last_synced else "—"
    msg_count = sum(len(v) for v in grouped.values())

    # Hero
    hero = f"""<div class="hero">
      <div class="hero-label">⬡ MsgStack &nbsp;·&nbsp; Message Spec</div>
      <h1>{spec.name}</h1>
      <p class="hero-tagline">{spec.tagline or spec.positioning[:100] if spec.positioning else ""}</p>
      {"<span class='hero-audience'>"+spec.audience+"</span>" if spec.audience else ""}
    </div>"""

    # Positioning card
    pos_card = f"""<div class="card">
      <div class="card-label">Positioning</div>
      <p class="positioning-text">{spec.positioning or "—"}</p>
      {"<p class='diff-text'>"+spec.differentiation+"</p>" if spec.differentiation else ""}
    </div>"""

    # Key messages grid
    section_order = ["Headline", "Subhead", "Benefit", "Use Case", "Proof Point", "QAPair", "Social Proof", "Positioning"]
    blocks = ""
    for sec in section_order:
        msgs = grouped.get(sec, [])
        if not msgs:
            continue
        meta = _SECTION_META.get(sec, {"icon": "◉", "color": "#475569", "bg": "#f8fafc"})
        items = "".join(
            f'<div class="msg-item"><span class="msg-num" style="background:{meta["bg"]};color:{meta["color"]}">{i+1}</span><span class="msg-text">{m}</span></div>'
            for i, m in enumerate(msgs[:4])
        )
        blocks += f"""<div class="section-block" style="background:{meta["bg"]}">
        <div class="section-title" style="color:{meta["color"]}"><span class="section-icon">{meta["icon"]}</span>{sec}</div>
        {items}
      </div>"""

    msgs_card = f"""<div class="card">
      <div class="card-label">Spec Entries &nbsp;·&nbsp; {msg_count} total</div>
      <div class="messages-grid">{blocks}</div>
    </div>"""

    # Personas
    audience_items = ""
    for p in audiences:
        initials = "".join(w[0].upper() for w in p.name.split()[:2])
        pain_tags = "".join(f'<span class="pain-tag">{pp[:40]}</span>' for pp in (p.pain_points or [])[:3])
        audience_items += f"""<div class="audience-card">
        <div class="audience-avatar">{initials}</div>
        <div class="audience-name">{p.name}</div>
        <div class="audience-desc">{(p.description or "")[:160]}</div>
        {pain_tags}
      </div>"""

    audiences_card = f"""<div class="card">
      <div class="card-label">Target Personas &nbsp;·&nbsp; {len(audiences)} defined</div>
      <div class="audiences-grid">{audience_items}</div>
    </div>"""

    footer = f"""<div class="footer">
      <span class="footer-badge">⬡ msgstack MCP &nbsp;·&nbsp; {last_synced} &nbsp;·&nbsp; {msg_count} messages &nbsp;·&nbsp; {len(audiences)} audiences</span>
    </div>"""

    body = hero + pos_card + msgs_card + audiences_card + footer
    return _base_html(f"{spec.name} — Messaging One Pager", body)


def _render_social_posts(spec, messages: list, channels: list) -> str:
    posts = []
    for m in messages:
        for ch in channels:
            variant = (m.variants or {}).get(ch)
            if variant:
                posts.append({
                    "channel": ch.title(),
                    "section": str(m.assertion_type).replace("_", " ").title(),
                    "content": variant,
                    "priority": m.priority,
                })

    ch_label = ", ".join(c.title() for c in channels)
    hero = f"""<div class="hero">
      <div class="hero-label">⬡ MsgStack &nbsp;·&nbsp; Social Posts</div>
      <h1>{spec.name}</h1>
      <p class="hero-tagline">{len(posts)} posts ready for {ch_label}</p>
      {"<span class='hero-audience'>"+spec.audience+"</span>" if spec.audience else ""}
    </div>"""

    if not posts:
        posts_html = '<div class="card"><p style="color:#64748b">No channel variants found. Add LinkedIn variants to your messages.</p></div>'
    else:
        posts_html = ""
        ch_colors = {"linkedin": ("#0a66c2", "#e8f0fe"), "twitter": ("#000", "#f0f0f0"), "email": ("#6366f1", "#eef2ff")}
        for post in posts:
            meta = _SECTION_META.get(post["section"], {"icon": "◉", "color": "#475569", "bg": "#f8fafc"})
            ch = post["channel"].lower()
            ch_color, ch_bg = ch_colors.get(ch, ("#475569", "#f8fafc"))
            posts_html += f"""<div class="post-card">
        <div>
          <span class="post-badge" style="background:{ch_bg};color:{ch_color}">{post['channel']}</span>
          <span class="post-badge" style="background:{meta['bg']};color:{meta['color']}">{meta['icon']} {post['section']}</span>
        </div>
        <p class="post-text">{post['content']}</p>
      </div>"""

    footer = '<div class="footer"><span class="footer-badge">⬡ msgstack MCP</span></div>'
    return _base_html(f"{spec.name} — Social Posts", hero + posts_html + footer)


def _render_email_template(spec, messages: list, stage: str) -> str:
    headlines = [m for m in messages if str(m.assertion_type) == "headline"]
    benefits  = [m for m in messages if str(m.assertion_type) == "benefit"]
    proofs    = [m for m in messages if str(m.assertion_type) == "proof_point"]

    stage_map = {
        "awareness": {
            "subject": headlines[0].content[:70] if headlines else (spec.tagline or ""),
            "hook": benefits[0].content if benefits else spec.positioning,
            "body": spec.differentiation or spec.positioning,
            "cta": f"See how {spec.name} works →",
        },
        "consideration": {
            "subject": f"How teams like yours use {spec.name}",
            "hook": proofs[0].content if proofs else (benefits[0].content if benefits else spec.positioning),
            "body": spec.positioning,
            "cta": "Book a 30-min demo →",
        },
        "decision": {
            "subject": f"Ready to get started with {spec.name}?",
            "hook": (benefits[0].variants or {}).get("email", benefits[0].content) if benefits else spec.positioning,
            "body": spec.differentiation or spec.positioning,
            "cta": "Start your free trial →",
        },
    }
    content = stage_map.get(stage, stage_map["awareness"])
    stages = ["awareness", "consideration", "decision"]
    tabs = "".join(
        f'<span class="stage-tab {"active" if s == stage else ""}">{s.title()}</span>'
        for s in stages
    )

    hero = f"""<div class="hero">
      <div class="hero-label">⬡ MsgStack &nbsp;·&nbsp; Email Template</div>
      <h1>{spec.name}</h1>
      <p class="hero-tagline">{stage.title()} stage · outbound email</p>
    </div>"""

    email_card = f"""<div class="card">
      <div class="card-label">Funnel Stage</div>
      <div class="stage-tabs">{tabs}</div>
      <div class="card-label" style="margin-top:8px">Email Content</div>
      <div class="email-field"><div class="email-label">Subject Line</div><div class="email-value subject">{content['subject']}</div></div>
      <div class="email-field"><div class="email-label">Opening Hook</div><div class="email-value">{content['hook']}</div></div>
      <div class="email-field"><div class="email-label">Body Copy</div><div class="email-value">{content['body']}</div></div>
      <div class="email-field"><div class="email-label">Call to Action</div><div class="email-value cta">{content['cta']}</div></div>
    </div>"""

    footer = '<div class="footer"><span class="footer-badge">⬡ msgstack MCP</span></div>'
    return _base_html(f"{spec.name} — Email ({stage.title()})", hero + email_card + footer)


def _render_battlecard(spec, messages: list, competitor: str) -> str:
    qa_pairs = [m for m in messages if str(m.assertion_type).endswith("qa_pair")]
    proofs = [m for m in messages if str(m.assertion_type).endswith("proof_point")]
    benefits = [m for m in messages if str(m.assertion_type).endswith("benefit")]

    extra = """
  .bc-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 20px; }
  @media (max-width: 640px) { .bc-grid { grid-template-columns: 1fr; } }
  .bc-col { border-radius: 12px; padding: 20px; }
  .bc-col.ours { background: #ecfdf5; border: 1.5px solid #34d399; }
  .bc-col.theirs { background: #fef2f2; border: 1.5px solid #f87171; }
  .bc-col-title { font-size: 12px; font-weight: 700; letter-spacing: .08em; text-transform: uppercase; margin-bottom: 12px; }
  .bc-col.ours .bc-col-title { color: #059669; }
  .bc-col.theirs .bc-col-title { color: #dc2626; }
  .bc-row { font-size: 13.5px; color: #334155; padding: 6px 0; border-bottom: 1px solid rgba(0,0,0,0.06); line-height: 1.5; }
  .bc-row:last-child { border-bottom: none; }
  .qa_pair-row { padding: 10px 12px; background: #fff; border: 1px solid #e2e8f0; border-radius: 8px; margin-bottom: 8px; }
  .qa_pair-label { font-size: 11px; font-weight: 600; text-transform: uppercase; letter-spacing: .06em; color: #94a3b8; margin-bottom: 4px; }
  .qa_pair-text { font-size: 13.5px; color: #1e293b; line-height: 1.5; }
"""
    hero = f"""<div class="hero">
      <div class="hero-label">⬡ MsgStack &nbsp;·&nbsp; Battlecard</div>
      <h1>{spec.name} vs {competitor}</h1>
      <p class="hero-tagline">{spec.tagline or spec.positioning[:80] if spec.positioning else ""}</p>
    </div>"""

    our_items = "".join(f'<div class="bc-row">✓ {b.content}</div>' for b in benefits[:5])
    their_items = "".join(f'<div class="bc-row">✗ {o.content}</div>' for o in qa_pairs[:4]) or \
        '<div class="bc-row" style="color:#94a3b8;">Add competitor weaknesses via qa_pair messages</div>'

    compare_card = f"""<div class="card">
      <div class="card-label">Head-to-Head Comparison</div>
      <div class="bc-grid">
        <div class="bc-col ours">
          <div class="bc-col-title">✦ {spec.name}</div>
          {our_items or '<div class="bc-row" style="color:#94a3b8;">Add benefit messages</div>'}
        </div>
        <div class="bc-col theirs">
          <div class="bc-col-title">✗ {competitor}</div>
          {their_items}
        </div>
      </div>
    </div>"""

    pos_card = f"""<div class="card">
      <div class="card-label">Positioning Against {competitor}</div>
      <p class="positioning-text">{spec.positioning or "—"}</p>
      {"<p class='diff-text'>" + spec.differentiation + "</p>" if spec.differentiation else ""}
    </div>"""

    proof_items = "".join(f'<div class="bc-row">◆ {p.content}</div>' for p in proofs[:5])
    proof_card = f"""<div class="card">
      <div class="card-label">Proof Points</div>
      {proof_items or '<p style="color:#94a3b8;font-size:13px;">Add proof_point messages to populate this section</p>'}
    </div>""" if proofs else ""

    obj_items = "".join(f'''<div class="qa_pair-row">
        <div class="qa_pair-label">QAPair {i+1}</div>
        <div class="qa_pair-text">{o.content}</div>
      </div>''' for i, o in enumerate(qa_pairs[:6]))
    obj_card = f"""<div class="card">
      <div class="card-label">QAPair Responses</div>
      {obj_items or '<p style="color:#94a3b8;font-size:13px;">Add qa_pair messages to populate this section</p>'}
    </div>"""

    footer = '<div class="footer"><span class="footer-badge">⬡ msgstack MCP · Battlecard</span></div>'
    return _base_html(f"{spec.name} vs {competitor} — Battlecard", hero + compare_card + pos_card + proof_card + obj_card + footer, extra)


def _render_email_sequence(spec, messages: list) -> str:
    headlines = [m for m in messages if str(m.assertion_type).endswith("headline")]
    benefits = [m for m in messages if str(m.assertion_type).endswith("benefit")]
    proofs = [m for m in messages if str(m.assertion_type).endswith("proof_point")]

    extra = """
  .seq-grid { display: grid; grid-template-columns: repeat(3, 1fr); gap: 20px; }
  @media (max-width: 700px) { .seq-grid { grid-template-columns: 1fr; } }
  .seq-card { border-radius: 14px; padding: 24px; border: 1.5px solid #e2e8f0; background: #fff; position: relative; }
  .seq-number { width: 32px; height: 32px; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-weight: 700; font-size: 13px; color: #fff; margin-bottom: 14px; }
  .seq-stage { font-size: 10px; font-weight: 700; letter-spacing: .1em; text-transform: uppercase; margin-bottom: 8px; }
  .seq-subject { font-size: 15px; font-weight: 600; color: #0f172a; line-height: 1.4; margin-bottom: 10px; }
  .seq-body { font-size: 13px; color: #475569; line-height: 1.6; margin-bottom: 12px; }
  .seq-cta { display: inline-block; font-size: 13px; font-weight: 600; color: #6366f1; text-decoration: none; border-bottom: 1px solid #6366f1; }
  .seq-connector { display: none; }
  @media (min-width: 700px) { .seq-connector { display: flex; align-items: center; justify-content: center; font-size: 18px; color: #cbd5e1; } }
"""
    stages = [
        {
            "num": "1", "color": "#6366f1", "stage": "Awareness",
            "subject": headlines[0].content[:70] if headlines else (spec.tagline or spec.name),
            "body": benefits[0].content if benefits else spec.positioning or "",
            "cta": f"Learn how {spec.name} works →",
        },
        {
            "num": "2", "color": "#0891b2", "stage": "Consideration",
            "subject": f"How teams like yours use {spec.name}",
            "body": proofs[0].content if proofs else (benefits[1].content if len(benefits) > 1 else spec.differentiation or spec.positioning or ""),
            "cta": "See the case study →",
        },
        {
            "num": "3", "color": "#059669", "stage": "Decision",
            "subject": f"Ready to get started with {spec.name}?",
            "body": spec.differentiation or spec.positioning or "",
            "cta": "Start your free trial →",
        },
    ]

    hero = f"""<div class="hero">
      <div class="hero-label">⬡ MsgStack &nbsp;·&nbsp; Email Sequence</div>
      <h1>{spec.name}</h1>
      <p class="hero-tagline">3-stage nurture sequence · Awareness → Consideration → Decision</p>
    </div>"""

    cards = ""
    for i, s in enumerate(stages):
        cards += f"""<div class="seq-card">
        <div class="seq-number" style="background:{s['color']}">{s['num']}</div>
        <div class="seq-stage" style="color:{s['color']}">{s['stage']}</div>
        <div class="seq-subject">{s['subject']}</div>
        <div class="seq-body">{s['body'][:300]}</div>
        <span class="seq-cta">{s['cta']}</span>
      </div>"""

    seq_card = f"""<div class="card">
      <div class="card-label">3-Email Nurture Sequence</div>
      <div class="seq-grid">{cards}</div>
    </div>"""

    footer = '<div class="footer"><span class="footer-badge">⬡ msgstack MCP · Email Sequence</span></div>'
    return _base_html(f"{spec.name} — Email Sequence", hero + seq_card + footer, extra)


@app.get("/api/preview/{skill_id}/{spec_id}")
def get_artifact_preview(skill_id: str, spec_id: str, request: Request):
    """Get Prefab preview HTML for an artifact."""
    from src.pipeline.generator import ArtifactGenerator
    from src.artifacts.prefab_generator import build_artifact_preview

    generator = ArtifactGenerator(store, skills)

    try:
        artifact = generator.generate(skill_id, spec_id, {})
        prefab_app = build_artifact_preview(skill_id, artifact.sections, artifact.spec_name, artifact.spec_id)
        
        # Check if it's an HTMX request
        if request.headers.get("HX-Request"):
            return HTMLResponse(content=prefab_app.html())
            
        return {"html": prefab_app.html()}
    except Exception as e:
        log.error("Preview generation failed: %s", e)
        raise HTTPException(500, str(e))


@app.get("/api/artifacts/{artifact_id}/image")
def get_artifact_image(artifact_id: str):
    """Generate a high-fidelity PNG image using Satori + resvg-python."""
    import subprocess
    from resvg_py import render_svg
    from src.pipeline.generator import ArtifactGenerator
    
    record = store.get_artifact(UUID(artifact_id))
    if not record:
        raise HTTPException(404, "Artifact not found")

    # Minimal Tailwind HTML for Satori
    # We'll use a simple card layout for now
    html_content = f"""
    <div style="display: flex; flex-direction: column; width: 100%; height: 100%; background-color: white; padding: 40px; font-family: 'WorkSans';">
        <div style="display: flex; align-items: center; margin-bottom: 20px;">
            <div style="width: 40px; height: 40px; background-color: #6366f1; border-radius: 8px; display: flex; align-items: center; justify-content: center; color: white; font-weight: 800; font-size: 20px; margin-right: 15px;">M</div>
            <div style="font-size: 24px; font-weight: 700; color: #1e293b;">{record['spec_name']}</div>
        </div>
        <div style="font-size: 48px; font-weight: 800; color: #0f172a; margin-bottom: 10px; line-height: 1.1;">{record['skill_id'].replace('_', ' ').title()}</div>
        <div style="font-size: 20px; color: #64748b; margin-bottom: 30px;">Generated via MsgStack MCP</div>
        <div style="flex: 1; display: flex; flex-direction: column; background-color: #f8fafc; border-radius: 12px; padding: 30px; border: 1px solid #e2e8f0;">
            <div style="font-size: 18px; color: #475569; line-height: 1.6;">{record['raw_content'][:500]}...</div>
        </div>
    </div>
    """

    try:
        # 1. Run Node.js bridge to get SVG
        cmd = ["node", "src/artifacts/render.js", html_content]
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        svg_data = result.stdout

        # 2. Rasterize SVG to PNG
        png_data = render_svg(svg_data)
        
        return HTMLResponse(content=png_data, media_type="image/png")
    except Exception as e:
        log.error("Image generation failed: %s", e)
        raise HTTPException(500, f"Failed to generate image: {str(e)}")

# --- Source Connections ---

@app.get("/api/connections")
def list_connections():
    return {"connections": store.list_connections()}


@app.get("/api/connections/google-drive/connect")
def gdrive_connect(folder_id: str = "", workspace_id: str = "default"):
    """Start the Google Drive OAuth2 flow. Redirects the browser to Google."""
    import json
    from fastapi.responses import RedirectResponse
    from src.sources.google_drive import GoogleDriveConnector

    if not folder_id:
        raise HTTPException(400, "folder_id is required")

    connector = GoogleDriveConnector()
    if not connector.client_id:
        raise HTTPException(500, "GOOGLE_CLIENT_ID not configured")

    state = json.dumps({"folder_id": folder_id, "workspace_id": workspace_id})
    import base64
    state_b64 = base64.urlsafe_b64encode(state.encode()).decode()
    url = connector.get_auth_url(state=state_b64)
    return RedirectResponse(url)


@app.get("/api/connections/google-drive/callback")
def gdrive_callback(code: str = "", state: str = "", error: str = ""):
    """OAuth2 callback — exchanges code for tokens and creates the connection."""
    import base64
    import json
    import threading
    from fastapi.responses import RedirectResponse
    from src.sources.google_drive import GoogleDriveConnector
    from src.sources.sync import get_sync_engine

    if error:
        return RedirectResponse(f"/?error={error}#connections")

    try:
        state_data = json.loads(base64.urlsafe_b64decode(state.encode()).decode())
    except Exception:
        raise HTTPException(400, "Invalid state parameter")

    folder_id = state_data.get("folder_id", "")
    workspace_id = state_data.get("workspace_id", "default")

    connector = GoogleDriveConnector()
    try:
        tokens = connector.exchange_code(code)
    except Exception as exc:
        log.error("Google OAuth token exchange failed: %s", exc)
        return RedirectResponse("/?error=token_exchange_failed#connections")

    access_token = tokens["access_token"]
    refresh_token = tokens.get("refresh_token", "")

    try:
        account_email = connector.get_account_email(access_token)
    except Exception as exc:
        log.warning("Could not fetch account email: %s", exc)
        account_email = ""

    try:
        folder_name = connector.get_folder_name(folder_id, access_token)
    except Exception as exc:
        log.warning("Could not fetch folder name: %s", exc)
        folder_name = folder_id

    try:
        page_token = connector.get_initial_page_token(access_token)
    except Exception as exc:
        log.warning("Could not fetch initial page token: %s", exc)
        page_token = "1"

    connection = store.create_connection(
        provider="google_drive",
        folder_id=folder_id,
        account_email=account_email,
        folder_name=folder_name,
        access_token=access_token,
        refresh_token=refresh_token,
        page_token=page_token,
        workspace_id=workspace_id,
    )

    # Run initial folder scan in a background thread so the redirect is immediate
    def _bg_initial_sync(conn_id: str):
        try:
            get_sync_engine().initial_sync(conn_id)
        except Exception as exc:
            log.error("Initial sync failed for %s: %s", conn_id, exc)

    thread = threading.Thread(target=_bg_initial_sync, args=(connection["id"],), daemon=True)
    thread.start()

    return RedirectResponse("/#connections")


@app.delete("/api/connections/{connection_id}")
def delete_connection(connection_id: str):
    if not store.delete_connection(connection_id):
        raise HTTPException(404, "Connection not found")
    return {"deleted": True}


@app.post("/api/connections/{connection_id}/sync")
def trigger_sync(connection_id: str):
    """Trigger an immediate manual sync for a connection."""
    import threading
    from src.sources.sync import get_sync_engine

    conn = store.get_connection(connection_id)
    if not conn:
        raise HTTPException(404, "Connection not found")

    def _bg(cid: str):
        try:
            get_sync_engine().sync_connection(cid)
        except Exception as exc:
            log.error("Manual sync failed for %s: %s", cid, exc)

    threading.Thread(target=_bg, args=(connection_id,), daemon=True).start()
    return {"status": "syncing"}


@app.get("/api/connections/{connection_id}/files")
def list_connection_files(connection_id: str):
    conn = store.get_connection(connection_id)
    if not conn:
        raise HTTPException(404, "Connection not found")
    return {"files": store.list_source_files(connection_id)}


@app.post("/api/connections/{connection_id}/source-files/{drive_file_id}/resync")
def resync_source_file(connection_id: str, drive_file_id: str):
    """Mark a source file for re-ingestion on next sync cycle."""
    conn = store.get_connection(connection_id)
    if not conn:
        raise HTTPException(404, "Connection not found")
    sf = store.get_source_file_by_drive_id(connection_id, drive_file_id)
    if not sf:
        raise HTTPException(404, "Source file not found")
    # Delete the existing spec so a fresh one is created on re-ingest
    if sf.get("spec_id"):
        try:
            store.delete_spec(UUID(sf["spec_id"]))
        except Exception:
            pass
    # Reset to error so the error-file retry loop picks it up on next sync
    store.upsert_source_file(
        connection_id=connection_id,
        drive_file_id=drive_file_id,
        file_name=sf["file_name"],
        mime_type=sf["mime_type"],
        drive_modified_at=sf.get("drive_modified_at", ""),
        sync_status="error",
        error_message="queued for resync",
    )
    return {"status": "queued"}


# --- Frontend ---

@app.get("/", response_class=HTMLResponse)
@app.get("/dashboard", response_class=HTMLResponse)
@app.get("/frameworks", response_class=HTMLResponse)
@app.get("/upload", response_class=HTMLResponse)
@app.get("/artifacts", response_class=HTMLResponse)
@app.get("/skills", response_class=HTMLResponse)
@app.get("/settings", response_class=HTMLResponse)
@app.get("/spec-detail", response_class=HTMLResponse)
@app.get("/connections", response_class=HTMLResponse)
def index(request: Request):
    return templates.TemplateResponse(request, "dashboard.html")


@app.get("/canvas", response_class=HTMLResponse)
def serve_canvas(request: Request):
    """Serve the Fabric.js canvas editor SPA."""
    return FileResponse("src/web/canvas/index.html")


# ── v0.7: Channel CRUD ────────────────────────────────────────────────────────

class ChannelCreate(BaseModel):
    name: str
    description: str = ""


class ChannelUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None


@app.get("/api/channels")
def list_channels():
    """List all channels with per-channel message counts."""
    channels = store.get_channels()
    for ch in channels:
        ch["message_count"] = store.get_channel_message_count(ch["id"])
    return channels


@app.get("/api/channels/{channel_id}")
def get_channel(channel_id: str):
    ch = store.get_channel(channel_id)
    if not ch:
        raise HTTPException(404, "Channel not found")
    ch["message_count"] = store.get_channel_message_count(channel_id)
    return ch


@app.post("/api/channels", status_code=201)
def create_channel(data: ChannelCreate):
    try:
        return store.create_channel(data.name, data.description)
    except ValueError as e:
        raise HTTPException(409, str(e))


@app.patch("/api/channels/{channel_id}")
def update_channel(channel_id: str, data: ChannelUpdate):
    try:
        result = store.update_channel(channel_id, data.name, data.description)
    except ValueError as e:
        raise HTTPException(403, str(e))
    if not result:
        raise HTTPException(404, "Channel not found")
    return result


@app.delete("/api/channels/{channel_id}", status_code=204)
def delete_channel(channel_id: str):
    try:
        deleted = store.delete_channel(channel_id)
    except ValueError as e:
        raise HTTPException(403, str(e))
    if not deleted:
        raise HTTPException(404, "Channel not found")
    return None


# ── v0.7: Bulk Message/Entry Status ───────────────────────────────────────────

class BulkStatusUpdate(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    entry_ids: list[str] = Field(validation_alias=AliasChoices('entry_ids', 'message_ids'), serialization_alias='entry_ids')
    status: str
    approved_by: Optional[str] = None

    @property
    def message_ids(self) -> list[str]:
        return self.entry_ids

    @message_ids.setter
    def message_ids(self, value: list[str]) -> None:
        self.entry_ids = value


@app.patch("/api/canon-domains/{domain_id}/entries/bulk-status")
@app.patch("/api/specs/{spec_id}/messages/bulk-status")
def bulk_update_message_status(data: BulkStatusUpdate, spec_id: Optional[str] = None, domain_id: Optional[str] = None, auth: AuthContext = Depends(require_write)):
    """Bulk approve/lock/flag multiple entries at once."""
    actual_id = domain_id or spec_id
    try:
        spec_uuid = UUID(actual_id)
    except Exception:
        raise HTTPException(400, "Invalid ID")
    spec = store.get_spec(spec_uuid)
    if not spec:
        raise HTTPException(404, "Spec domain not found")
    if not auth.has_department_access(spec.department):
        raise HTTPException(403, f"You do not have permission to write to this domain's department.")
    try:
        count = store.bulk_update_message_status(
            data.message_ids, data.status, data.approved_by or "admin"
        )
    except ValueError as e:
        raise HTTPException(400, str(e))
    return {"ok": True, "updated_count": count}


# ── v0.7: Review Log ──────────────────────────────────────────────────────────

@app.get("/api/canon-domains/{domain_id}/review-log")
@app.get("/api/specs/{spec_id}/review-log")
def get_message_review_log(limit: int = Query(50, ge=1, le=200), spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    """Return review log for a spec (approvals, locks, reviews)."""
    actual_id = domain_id or spec_id
    try:
        spec_uuid = UUID(actual_id)
    except Exception:
        raise HTTPException(400, "Invalid ID")
    if not store.get_spec(spec_uuid):
        raise HTTPException(404, "Spec domain not found")
    return store.get_review_log(actual_id, limit=limit)


# ── v0.7: Mark Spec Domain Reviewed ──────────────────────────────────────────

@app.post("/api/canon-domains/{domain_id}/mark-reviewed")
@app.post("/api/specs/{spec_id}/mark-reviewed")
def mark_spec_reviewed_v2(reviewed_by: Optional[str] = Query(None), spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    """Mark a domain as reviewed (updates last_reviewed, appends review log)."""
    actual_id = domain_id or spec_id
    try:
        spec_uuid = UUID(actual_id)
    except Exception:
        raise HTTPException(400, "Invalid ID")
    result = store.mark_spec_reviewed(str(spec_uuid), reviewed_by or "admin")
    if not result:
        raise HTTPException(404, "Spec domain not found")
    return {"ok": True, **result}


# ── v0.7: Artifact Ratings & Feedback Loop ───────────────────────────────────

class ArtifactRatingCreate(BaseModel):
    rating: int              # 1-5
    tag: str = "good"        # "good" | "bad"
    rated_by: Optional[str] = None
    notes: Optional[str] = None


@app.post("/api/artifacts/{artifact_id}/rate", status_code=201)
def rate_artifact(artifact_id: str, data: ArtifactRatingCreate):
    """Rate a generated artifact (1-5 stars). Updates chunk boost factors."""
    try:
        result = store.record_artifact_rating(
            artifact_id=artifact_id,
            rating=data.rating,
            tag=data.tag,
            rated_by=data.rated_by or "",
            notes=data.notes or "",
        )
    except ValueError as e:
        raise HTTPException(400, str(e))
    return result


@app.get("/api/canon-domains/{domain_id}/usage-stats")
@app.get("/api/specs/{spec_id}/usage-stats")
def get_spec_usage_stats(spec_id: Optional[str] = None, domain_id: Optional[str] = None):
    """Return assertion usage heatmap — times used and avg rating, sorted by usage."""
    actual_id = domain_id or spec_id
    try:
        spec_uuid = UUID(actual_id)
    except Exception:
        raise HTTPException(400, "Invalid ID")
    if not store.get_spec(spec_uuid):
        raise HTTPException(404, "Spec domain not found")
    return store.get_message_usage_stats(actual_id)


# --- Spec Navigator Chat ---

class ChatRequest(BaseModel):
    query: str
    workspace_id: str = "default"


@app.post("/api/chat")
def api_chat_stream(req: ChatRequest, auth: AuthContext = Depends(get_auth_context)):
    """SSE streaming endpoint for the Spec Navigator conversational chat."""
    try:
        from src.pipeline.agents import SpecNavigator
        from openai import OpenAI
        import os as _os

        # Query audit (non-blocking) — chat is the web-surface query path
        try:
            from src.models import QueryAuditLog
            store.log_query(QueryAuditLog(
                workspace_id=req.workspace_id or "default",
                user_id=auth.name or "web",
                query_text=req.query,
                source="web:chat",
            ))
        except Exception as _log_exc:
            log.warning("Query audit logging failed (non-blocking): %s", _log_exc)

        client = OpenAI(api_key=_os.environ.get("OPENAI_API_KEY"))
        navigator = SpecNavigator(client, store)

        def event_generator():
            for chunk in navigator.chat_stream(req.query, workspace_id=req.workspace_id):
                yield f"data: {json.dumps({'content': chunk})}\n\n"
            yield "data: [DONE]\n\n"

        return StreamingResponse(event_generator(), media_type="text/event-stream")
    except Exception as e:
        raise HTTPException(500, f"Streaming connection failed: {e}")


@app.get("/{full_path:path}", response_class=HTMLResponse)
def catch_all(request: Request, full_path: str):
    """Serve the SPA for any unmatched path so page refreshes don't 404."""
    return templates.TemplateResponse(request, "dashboard.html")


# --- Penpot Webhook ---

@app.post("/api/webhooks/penpot")
async def penpot_webhook(request: Request):
    """Handle incoming Penpot webhooks for design changes.

    Updates MsgStack brand settings when Penpot designs change.
    """
    try:
        body = await request.json()
    except Exception:
        raise HTTPException(400, "Invalid JSON body")

    from src.design.penpot_sync import handle_penpot_webhook
    result = handle_penpot_webhook(body)
    return result


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
